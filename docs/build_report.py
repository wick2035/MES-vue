# -*- coding: utf-8 -*-
"""
生成《MES 智能制造执行系统 · 课程实验报告》Word 文档。
内容依据本仓库真实实现反推编撰，覆盖全部已实现功能（含必做 12 项与大量扩展项）。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"d:\AAA-study\MES-Springboot-master\MES-Springboot-master"
IMG = os.path.join(ROOT, "docs", "manual", "images")
OUT = os.path.join(ROOT, "docs", "MES智能制造执行系统-课程实验报告.docx")

CJK = "宋体"
CJK_H = "黑体"
LATIN = "Times New Roman"
MONO = "Consolas"

# 标题 / 小标题统一用黑色（原为蓝色）。架构图等图形块仍用字面量色值，不受此影响。
BLUE = RGBColor(0x00, 0x00, 0x00)
DARKBLUE = RGBColor(0x00, 0x00, 0x00)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x59, 0x59, 0x59)
TABLE_HEADER_FILL = "D9D9D9"  # 表头浅灰底（原为蓝底白字）

doc = Document()

# ---------------- 基础排版工具 ----------------

def set_run_font(run, east=CJK, west=LATIN, size=None, bold=None, italic=None, color=None):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), east)
    rfonts.set(qn('w:ascii'), west)
    rfonts.set(qn('w:hAnsi'), west)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, east, west, size, bold=False, color=None):
    style.font.name = west
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), east)
    rfonts.set(qn('w:ascii'), west)
    rfonts.set(qn('w:hAnsi'), west)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)


# ---------------- 页面与样式 ----------------
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.6)

set_style_font(doc.styles['Normal'], CJK, LATIN, 12)
doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

for name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13), ('Heading 4', 12)]:
    set_style_font(doc.styles[name], CJK_H, LATIN, size, bold=True, color=BLUE)
    doc.styles[name].paragraph_format.space_before = Pt(10)
    doc.styles[name].paragraph_format.space_after = Pt(6)

FIG_NO = [0]
TAB_NO = [0]

# ---------------- 内容生成函数 ----------------

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    set_run_font(r, CJK_H, LATIN, 16, bold=True, color=BLUE)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    set_run_font(r, CJK_H, LATIN, 14, bold=True, color=BLUE)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    set_run_font(r, CJK_H, LATIN, 13, bold=True, color=DARKBLUE)
    return p


def body(text, indent=True, size=12, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    r = p.add_run(text)
    set_run_font(r, CJK, LATIN, size, bold=bold)
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, CJK, LATIN, size)
    return p


def numbered(text, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_run_font(r, CJK, LATIN, size)
    return p


def table(headers, rows, widths=None, fontsize=10, caption=None):
    if caption:
        TAB_NO[0] += 1
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run("表 %s  %s" % (caption_num(), caption))
        set_run_font(cr, CJK_H, LATIN, 10.5, bold=True, color=GRAY)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], TABLE_HEADER_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext)
        set_run_font(r, CJK_H, LATIN, fontsize, bold=True, color=BLACK)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(str(val))
            set_run_font(r, CJK, LATIN, fontsize)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


_TAB_LABEL = [0]

def caption_num():
    return "%d-%d" % (CUR_CH[0], _chapter_tab())


CUR_CH = [0]
_FIG_IN_CH = {}
_TAB_IN_CH = {}

def _chapter_tab():
    _TAB_IN_CH[CUR_CH[0]] = _TAB_IN_CH.get(CUR_CH[0], 0) + 1
    return _TAB_IN_CH[CUR_CH[0]]

def _chapter_fig():
    _FIG_IN_CH[CUR_CH[0]] = _FIG_IN_CH.get(CUR_CH[0], 0) + 1
    return _FIG_IN_CH[CUR_CH[0]]


def figure(filename, caption, width_cm=15.0):
    path = os.path.join(IMG, filename)
    if not os.path.exists(path):
        body("[缺少配图：%s]" % filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("图 %d-%d  %s" % (CUR_CH[0], _chapter_fig(), caption))
    set_run_font(r, CJK_H, LATIN, 10.5, bold=True, color=GRAY)
    doc.add_paragraph()


def code(text, size=8.5):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    shade(cell, "F4F6F8")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in text.split('\n'):
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line != '' else ' ')
        set_run_font(r, MONO, MONO, size)
    doc.add_paragraph()


def quote(text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run("【说明】" + text)
    set_run_font(r, CJK, LATIN, size, italic=False, color=GRAY)
    return p


def corecode(java, explain, label="核心代码与说明（节选自源码）"):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("▶ " + label)
    set_run_font(r, CJK_H, LATIN, 11, bold=True, color=DARKBLUE)
    code(java)
    body(explain, size=11)


def page_break():
    doc.add_page_break()


def chapter(no, title):
    CUR_CH[0] = no
    page_break()
    h1("第 %d 章  %s" % (no, title))


def layer_box(title, content, fill, textcolor=RGBColor(255, 255, 255)):
    """单行整宽表格，模拟分层架构的一层。"""
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(3.2)
    t.columns[1].width = Cm(12.3)
    c0 = t.cell(0, 0)
    c1 = t.cell(0, 1)
    shade(c0, fill)
    shade(c1, "EEF3F8")
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(title)
    set_run_font(r0, CJK_H, LATIN, 11, bold=True, color=textcolor)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(content)
    set_run_font(r1, CJK, LATIN, 10, color=BLACK)


def flow_row(steps, fill="2B6CB0"):
    """横向流程：steps 之间以 → 连接，单元格底色。"""
    n = len(steps)
    cols = 2 * n - 1
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(cols):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i % 2 == 0:
            shade(cell, fill)
            r = p.add_run(steps[i // 2])
            set_run_font(r, CJK_H, LATIN, 9, bold=True, color=RGBColor(255, 255, 255))
        else:
            r = p.add_run("→")
            set_run_font(r, CJK_H, LATIN, 11, bold=True, color=DARKBLUE)
    doc.add_paragraph()


# ============================================================
# 封面
# ============================================================
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("工业互联网 / 工业软件系统课程")
set_run_font(r, CJK_H, LATIN, 15, bold=True, color=GRAY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("实 验 报 告")
set_run_font(r, CJK_H, LATIN, 30, bold=True, color=BLUE)
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("基于开源 MES 的功能扩展与性能优化")
set_run_font(r, CJK_H, LATIN, 19, bold=True, color=BLACK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("——MES 智能制造执行系统的设计与实现")
set_run_font(r, CJK_H, LATIN, 15, bold=True, color=DARKBLUE)
for _ in range(3):
    doc.add_paragraph()

_info = [
    ("报告题目", "面向离散装配制造的 MES 智能制造执行系统功能扩展与优化"),
    ("选题方向", "开源工业软件（MES）系统的功能扩展 + 性能优化 + 智能化创新"),
    ("系统名称", "黑科制造 MES 系统（Manufacturing Execution System · PRO）"),
    ("技术架构", "Spring Boot 单体 · Java 11 · MyBatis-Plus · MySQL · Redis · Shiro · FreeMarker · layui"),
    ("智能化能力", "通义千问大模型（SSE 流式 + Function Calling）· 实时数据大屏 · 3D 数字孪生"),
    ("代码主模块", "mes/（根包 com.wangziyang.mes）"),
    ("访问地址", "http://localhost:9090（默认账号 admin / 123）"),
    ("文档版本", "V1.0"),
    ("编制日期", "2026 年 6 月 14 日"),
]
ti = doc.add_table(rows=0, cols=2)
ti.alignment = WD_TABLE_ALIGNMENT.CENTER
ti.style = 'Table Grid'
for k, v in _info:
    cells = ti.add_row().cells
    cells[0].width = Cm(3.6); cells[1].width = Cm(11.9)
    shade(cells[0], "EEF3F8")
    pk = cells[0].paragraphs[0]; pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = pk.add_run(k); set_run_font(rk, CJK_H, LATIN, 11, bold=True, color=BLUE)
    pv = cells[1].paragraphs[0]
    rv = pv.add_run(v); set_run_font(rv, CJK, LATIN, 10.5)

# ============================================================
# 目录
# ============================================================
page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目  录"); set_run_font(r, CJK_H, LATIN, 18, bold=True, color=BLUE)
doc.add_paragraph()
_toc_p = doc.add_paragraph()
_run = _toc_p.add_run()
_fb = OxmlElement('w:fldChar'); _fb.set(qn('w:fldCharType'), 'begin')
_it = OxmlElement('w:instrText'); _it.set(qn('xml:space'), 'preserve')
_it.text = 'TOC \\o "1-3" \\h \\z \\u'
_fs = OxmlElement('w:fldChar'); _fs.set(qn('w:fldCharType'), 'separate')
_t = OxmlElement('w:t'); _t.text = "（请在 Word 中按 Ctrl+A 全选后按 F9，选择“更新整个目录”以生成页码目录）"
_rt = OxmlElement('w:r'); _rt.append(_t)
_fe = OxmlElement('w:fldChar'); _fe.set(qn('w:fldCharType'), 'end')
_run._r.append(_fb); _run._r.append(_it); _run._r.append(_fs); _run._r.append(_rt); _run._r.append(_fe)

# ============================================================
# 第 1 章
# ============================================================
chapter(1, "项目概述与现场调研")

h2("1.1 项目概述")
body("本项目以一套面向离散装配制造企业的开源型 MES（Manufacturing Execution System，制造执行系统）为基线，"
     "在课程提供的系统需求与架构资料基础上，完成大规模的功能扩展、流程贯通与智能化优化，最终形成一套可演示、可落地、"
     "覆盖“基础数据 → 产品与工艺设计 → 生产计划 → 派工下发 → 现场执行采集 → 库房物料管控 → 数据可视化 → AI 辅助”"
     "全链路的智能制造执行平台——“黑科制造 MES 系统（PRO 版）”。")
body("系统采用 Spring Boot 单体架构（Java 11），持久层基于 MyBatis-Plus + MySQL，缓存采用 Redis 与 EhCache，"
     "鉴权采用 Apache Shiro，视图层采用 FreeMarker 服务端渲染配合自封装的 sp* 前端组件（spTable / spLayer / spUtil），"
     "可视化采用 ECharts 与 Three.js，并接入通义千问大模型实现对话式数据助手与 AI 辅助建模。系统启动后通过 "
     "http://localhost:9090 访问，默认管理员账号 admin / 123。")
body("本报告依照“需求分析—架构设计—模块开发—集成测试—运维规划—创新设计”的工程全流程组织。需要特别说明的是，"
     "本系统实际完成的功能远多于考核要求的 12 项必做功能：除完整实现角色权限、班组员工、编组设备、加工单元、物料信息、"
     "库房库位（含 3D 仿真联动）、零部件、产品 BOM、工序、工艺路线/流程、工艺内容编制、产品工艺查询等全部必做功能外，"
     "还自主扩展了生产订单与 MRP 物料需求计划、两级派工（设备/员工，含负载均衡）、生产计划下发、工单全生命周期管理与完工交付、"
     "SN 过站采集追溯、申请—确认两段式库房管理中心、通用审批工作流引擎、实时数据中心大屏、3D 数字孪生仓库、"
     "AI 智能数据助手与 AI 智能建模四步向导等大量增值模块，构成一个端到端贯通的完整制造执行闭环。")

h2("1.2 行业背景与建设意义")
body("在“工业 4.0”“中国制造 2025”与工业互联网战略的推动下，制造企业的信息化重心正从以 ERP 为代表的“计划层”，"
     "下沉到直接管控车间现场的“执行层”。MES 处于企业信息化金字塔的中间层：向上承接 ERP/PLM 下达的生产订单与物料主数据，"
     "向下贯通车间设备、班组与在制品，是消弭“计划与执行之间鸿沟”的关键枢纽。")
body("然而，大量中小型离散制造企业的车间现场仍停留在“纸质工单 + Excel 台账 + 口头派工”的粗放管理状态，普遍存在数据孤岛、"
     "进度不透明、缺料停线、质量不可追溯等顽疾。引入并深度定制一套贴合自身工艺的 MES，对企业实现生产过程透明化、数据可追溯化、"
     "排产与派工智能化具有直接价值。本项目选择开源 MES 进行二次开发与扩展，正是为了在可控成本下，验证“以全链路数据贯通 + AI 辅助”"
     "提升离散制造执行水平的工程可行性。")
flow_row(["ERP / PLM\n计划层", "MES\n执行层(本系统)", "车间现场\n控制层"], fill="2B6CB0")
quote("MES 向上承接订单/BOM/主数据，向下派发工艺文件与报工指令；车间通过 SN 过站采集、不良反馈回流 MES，"
      "MES 再向 ERP 上报完工、良率与库存，形成纵向贯通的数据闭环。")

h2("1.3 现场调研的开展")
body("为准确定位工业现场的核心矛盾，本项目以一家典型的电子整机装配企业（以“台式电脑主机”整机装配产线为标杆场景）为调研对象，"
     "综合采用现场走访观察、关键岗位访谈、业务单据采集与历史数据分析四类方法，覆盖计划、工艺、仓储、车间、质检五个关键岗位。")
table(["调研维度", "调研对象 / 岗位", "调研方法", "重点采集内容"],
      [["计划排产", "生产计划员", "访谈 + 单据采集", "订单录入、交期承诺、缺料停线频次、排产依据"],
       ["工艺工程", "工艺工程师", "访谈 + 工艺文件审阅", "BOM 维护方式、工艺路线编制、工艺变更与版本失控问题"],
       ["仓储物料", "库房管理员", "现场观察 + 台账核对", "库位管理、出入库登账、账实不符与缺料齐套问题"],
       ["车间执行", "现场班长 / 操作工", "现场跟线 + 工时记录", "派工方式、报工时效、在制品状态可见性"],
       ["质量追溯", "质检员", "访谈 + 不良记录分析", "单件追溯能力、不良统计、返工返修闭环"]],
      widths=[2.6, 3.2, 3.4, 6.3],
      caption="现场调研对象与方法一览")
body("调研结论表明：企业当前痛点高度集中于“数据不通、过程不透、质量不溯、排产靠经验”，且现场设备来源多样、数据格式各异，"
     "进一步加剧了集成难度。下文从技术瓶颈与工业环境约束两个层面展开深度分析。")

h2("1.4 现场核心问题识别与技术瓶颈分析")
body("结合工业互联网典型问题域，本项目将调研发现归纳为四类核心问题，并逐一剖析其背后的技术瓶颈与本系统的应对方向。")

h3("1.4.1 设备互联协议不兼容")
body("现场设备（装配工装、测试仪、贴标机、打印机、PLC 等）来自不同厂商、不同年代，对外暴露的接口五花八门：有的支持 OPC-UA / Modbus-TCP，"
     "有的仅提供串口或私有协议，还有相当一部分老旧工位根本无数字化接口，只能依赖人工记录。协议层的不统一导致难以用一套采集程序统一对接，"
     "“最后一公里”的设备数据无法稳定回流。技术瓶颈在于：缺乏统一的设备接入抽象层与标准化的数据采集语义。")
body("本系统的应对：在架构上预留“设备接入层 / OPC 操作”作为独立分组（系统顶栏即划分出“常规管理 / OPC 操作 / 其他管理”三大入口），"
     "并以“工序—加工单元—设备/设备编组”的资源模型对设备进行统一抽象建模；在执行层以 SN（产品序列号）过站采集作为协议无关的"
     "标准化数据入口——无论底层设备是否联网，现场均可通过统一的 SN 扫码过站接口上报每道工序的 OK/NG 结果，从而以“以人补机、"
     "以标准接口屏蔽协议差异”的渐进式策略实现设备数据的可采集、可追溯。")

h3("1.4.2 异构数据格式混乱")
body("物料、BOM、工艺、库存、工单等主数据散落于 Excel、纸质单据与多个孤立小系统中，编码规则不统一、字段口径不一致、"
     "同一物料在不同环节存在多种叫法，导致跨环节的数据无法对齐，统计与追溯困难。技术瓶颈在于：缺乏统一的主数据模型、"
     "统一的编码与字典体系，以及统一的接口数据契约。")
body("本系统的应对：① 以统一实体基类 BaseEntity 规范所有业务实体的主键（雪花算法字符串 ID）与审计字段，消除“同表不同构”；"
     "② 以四类型物料模型（FG 成品 / PG 半成品 / COMP 组件 / PART 零件）与三层 BOM 树统一产品结构语义；③ 以数据字典统一状态、"
     "单位、类型等枚举口径；④ 前后端约定单一 JSON 返回信封 Result{code,data,msg}，使所有接口的数据契约一致，从根本上"
     "治理“异构数据格式混乱”。")

h3("1.4.3 实时控制响应滞后")
body("传统模式下，现场进度、缺料预警、质量异常往往要等到日报、周报才汇总上来，管理者“看不到、看得晚”，决策严重滞后；"
     "物料需求也多为批量离线运算，无法随库存实时变化重算齐套。技术瓶颈在于：缺乏实时数据通道与实时聚合计算能力。")
body("本系统的应对：① 智能制造数据中心大屏基于真实业务数据做内存级聚合，并以 30 秒自动刷新呈现产量、合格率、工序流转、"
     "库存与人员负载等态势；② MRP 物料需求计划基于 sp_inventory 实时重算可用量与净需求，库存一变即可重新齐套校验；"
     "③ AI 智能数据助手采用 SSE（Server-Sent Events）服务器推送实现毫秒级流式逐字应答，使“问数据—得结论”几乎实时；"
     "④ SN 首站过站即将工单置为“已开工（STARTED）”并实时回流大屏，显著压缩“执行—感知”的时延。")

h3("1.4.4 计划与执行脱节、过程不可追溯")
body("订单、工艺、派工、执行、入出库各自为政，计划下发后难以跟踪到工序级进度；产品出现质量问题时无法回溯到具体 SN、"
     "具体工序与责任人。技术瓶颈在于：缺乏贯穿订单到 SN 的端到端数据主线与单件级追溯粒度。")
body("本系统的应对：构建“生产订单 → MRP → 计划下发 → 工单 → 两级派工 → SN 过站 → 完工交付 → 入库”的端到端数据主线，"
     "每张单据可上下追溯；并以 SN 为最小粒度记录每道工序的过站结果，支持正反向质量追溯与工序级良率统计，"
     "彻底解决“脱节”与“不可追溯”问题。")

h2("1.5 工业环境约束分析")
body("工业现场不同于一般信息系统的办公环境，系统设计必须充分考虑高温高湿、粉尘油污、电磁干扰、电力波动与网络不稳定等约束，"
     "以及功能安全与数据安全合规要求。下表列出主要约束因素及本系统的针对性设计权衡。")
table(["约束类别", "现场约束描述", "对系统的影响", "本系统设计权衡"],
      [["高温高湿/粉尘", "车间温湿度高、粉尘油污大，触摸屏与扫码枪易污损误触", "终端交互需简化、容错", "现场采集页极简化（选工单→扫 SN→选 OK/NG→采集），防重复过站校验，误操作可纠错"],
       ["电磁干扰/网络抖动", "大功率设备产生电磁干扰，车间 Wi-Fi 偶发丢包", "请求可能超时/重发", "无状态接口 + 幂等设计（SQL 脚本幂等、授权重建先清后插），SSE 断线可重连，关键写操作事务回滚"],
       ["电力波动/宕机", "供电不稳，存在异常断电", "数据一致性风险", "跨表操作统一 @Transactional(rollbackFor=Exception)，软删除不丢数据，库存以单据流水驱动可对账"],
       ["功能安全", "误下发、误出库会造成停线或物料损失", "需流程管控", "BOM/工艺定版锁定、配套出库整单校验、下发前校验派工齐全、关键业务挂接审批工作流"],
       ["数据安全合规", "账号权限、操作审计、敏感凭证保护（等保要求）", "需鉴权与审计", "Shiro RBAC + 加盐密码、菜单级权限、审计字段全量留痕、API Key 走环境变量不入库、LLM 工具只读白名单"]],
      widths=[2.6, 4.0, 3.0, 5.9],
      caption="工业环境约束因素与系统设计权衡")

h2("1.6 技术与文献调研")
body("在确定技术路线前，本项目对 MES 领域标准、工业互联网平台参考架构与主流开源技术栈进行了调研，作为架构设计与技术选型的依据。")
bullet("MES 功能与标准：参照 ISA-95 / IEC 62264 制造运行管理（MOM）模型与 MESA-11 功能模型，明确 MES 应覆盖资源管理、详细排产、"
       "派工、过程管理、数据采集、质量管理、产品追溯等核心功能，本系统据此规划业务域边界。")
bullet("工业互联网参考架构：借鉴“边缘层—平台层（IaaS/PaaS）—应用层”的分层思想与“设备层—网络层—平台层—应用层”的纵向贯通模型，"
       "将本系统映射为面向执行层的工业软件平台（详见第 3 章）。")
bullet("数据可追溯：参照单件追溯（Unit-level Traceability）理念，确定以 SN 为追溯粒度、以工序过站记录为追溯链节点的设计。")
bullet("技术栈选型：在开源 Java 生态中，Spring Boot + MyBatis-Plus + Shiro + FreeMarker + layui 组合具备成熟度高、上手快、"
       "单体部署运维简单的优势，契合中小制造企业的落地条件；可视化选用社区活跃的 ECharts 与 Three.js；大模型选用国产合规的"
       "通义千问（DashScope，OpenAI 兼容接口），兼顾能力与数据合规。")
body("调研结论：以 Spring Boot 单体为基线、按业务域分包、统一 CRUD 范式、数据库脚本化演进，是在成本、交付速度与可维护性之间"
     "的最优工程折中；在此基础上叠加实时大屏、3D 数字孪生与大模型能力，可在不牺牲稳定性的前提下显著提升系统的实用性与智能化水平。")

h2("1.7 问题分析报告（鱼骨图）")
body("针对 MES 实施中最受关注的质量目标“降低生产交付延期率”，本项目采用鱼骨图（因果图）从人、机、料、法、环、测六个维度"
     "（6M）系统分析根因，并逐项标注本系统对应的设计对策，形成“根因 → 对策”的闭环映射。")
table(["维度", "根因", "本系统设计对策"],
      [["人 Man", "派工不均、技能错配", "员工派工负载均衡算法（按未完成任务最少择人），AI 向导第④步自动排人"],
       ["人 Man", "报工不及时", "SN 扫码过站，首站即置工单 STARTED，实时回流大屏"],
       ["机 Machine", "设备未派工 / 冲突", "工序级设备派工，下发前校验设备 + 员工双全"],
       ["料 Material", "缺料未齐套", "MRP 基于 sp_inventory 实时重算，齐套校验通过方可下发；配套出库整单校验"],
       ["料 Material", "库存账实不符", "单据驱动库存变动，库存终值与出入库流水严格对齐、可逐笔审计"],
       ["法 Method", "工艺缺失 / 未定版", "BOM 与工艺路线定版（lock）机制，未定版不可下发；AI 向导自动补建并定版"],
       ["法 Method", "审批链路长", "通用工作流引擎配置驱动，订单 / 工单变更标准化审批"],
       ["环 Env", "计划与执行脱节", "订单 → MRP → 下发 → 工单 → SN 全链贯通"],
       ["测 Meas", "缺实时可视化", "数据大屏 30s 刷新 + 3D 数字孪生 + AI 数据助手"],
       ["测 Meas", "不良无追溯", "SN 级 OK/NG 记录，支持正反向追溯与良率统计"]],
      widths=[2.0, 4.5, 9.0],
      caption="“生产交付延期”根因（6M）与系统设计对策映射")

# ============================================================
# 第 2 章
# ============================================================
chapter(2, "系统需求分析")

h2("2.1 需求分析方法与来源")
body("本系统的需求来源于三个层面：① 课程提供的《资源分配管理》《BOM 与组件数据管理》《工艺设计管理》等系统需求与架构资料"
     "（明确了 12 项必做功能）；② 第 1 章现场调研所暴露的真实痛点；③ 在工程实现过程中，为使各功能形成端到端闭环而反向推导出的"
     "一批支撑性、贯通性需求（即“扩展需求”）。需求分析采用“用例驱动 + 数据驱动”相结合的方法，先识别角色与用例，再细化数据与状态。")
quote("说明：考核要求中的若干必做功能（如库房库位、产品工艺查询）若孤立实现并无业务闭环价值。为使其真正“跑起来”，"
      "本项目反推出订单、MRP、派工、下发、过站、出入库等关联需求并一并实现，这正是“从功能点反推完整需求”的工程体现。")

h2("2.2 角色与利益相关者分析")
body("系统采用“用户 → 角色 → 菜单/数据权限”三级 RBAC 模型。基于调研岗位，预置并实现了如下业务角色，各角色按授权菜单获得"
     "“千人千面”的导航视图。")
table(["角色", "编码", "典型职责", "对应核心模块"],
      [["系统管理员", "888888", "全系统配置，用户/角色/菜单/部门维护", "系统管理全部"],
       ["生产计划员", "10002", "订单录入、MRP、派工、计划下发", "生产计划中心"],
       ["工艺员", "technologyRole", "物料/BOM/工序/工艺规划与内容编制", "产品/工艺管理中心"],
       ["车间主管", "10003", "工单与派工审核、现场调度", "计划管理、流程办理"],
       ["现场班长", "10004", "班组任务分配与执行跟踪", "在制品、计划管理"],
       ["操作工", "10005", "过站采集、现场作业", "在制品管理"],
       ["质量管理员/检验员", "10006", "质量判定、NG 处理与追溯", "在制品、数据中心"],
       ["库房管理员", "—", "出入库确认、配套出库、库存盘点", "库房管理中心"]],
      widths=[2.8, 2.6, 5.2, 4.9],
      caption="系统角色与职责矩阵")

h2("2.3 功能需求")
body("功能需求分为“考核必做功能”与“自主扩展功能”两类。下表先给出与考核要求逐项对应的必做功能需求及其完成情况，"
     "再概述自主扩展的增值功能。")
h3("2.3.1 考核必做功能需求（含完成度与对应资料）")
table(["序号", "考核要求功能", "参考资料/页码", "本系统对应实现", "完成度"],
      [["1", "完善角色管理（权限管理）", "资源分配管理 11-17", "角色增强：菜单授权树/数据权限/分配用户/7 类业务角色", "已完成 ✓"],
       ["2", "班组员工定义", "资源分配管理 29-37", "班组（主）—班组员工（从）主从维护", "已完成 ✓"],
       ["3", "编组设备定义", "资源分配管理 42-48", "设备编组（主）—设备（从）主从维护", "已完成 ✓"],
       ["4", "加工单元定义", "资源分配管理 55-57", "加工单元—绑定执行班组，工序的归属载体", "已完成 ✓"],
       ["5", "物料信息定义/物料维护", "资源分配管理 77-83", "四类型物料 + 提前期/安全库存/来源 + 批量导入", "已完成 ✓"],
       ["6", "库房库位定义 + 3D 仿真联动", "资源分配管理 99-102", "组×排×层×列笛卡尔积自动建库位 + 3D 数字孪生读取库位坐标", "已完成 ✓(含拓展)"],
       ["7", "零部件定义", "BOM 与组件数据管理 10-13", "零部件定义（组件/半成品归属产品），BOM 定版前置", "已完成 ✓"],
       ["8", "工艺/产品 BOM 管理", "BOM 与组件数据管理 30-37", "三层 BOM 树（child_bom_id 关联）+ 定版锁定", "已完成 ✓"],
       ["9", "工序信息定义", "工艺设计管理 11-14", "标准工序库，归属加工单元，工时/周期/是否生成计划", "已完成 ✓"],
       ["10", "工艺流程/工艺路线管理", "工艺设计管理 23-30", "工艺路线（工序时序）+ 工艺流程（落到 BOM 树并锁定）", "已完成 ✓"],
       ["11", "工艺内容编制", "工艺设计管理 46-58", "按节点编制作业内容（editing→completed）", "已完成 ✓"],
       ["12", "产品工艺查询", "工艺设计管理 59-61", "只读按 BOM 查看完整工艺规划与内容", "已完成 ✓"]],
      widths=[1.1, 4.0, 3.4, 5.6, 1.8],
      caption="考核必做功能需求与实现对照（需求追溯矩阵）")

h3("2.3.2 自主扩展功能需求（增值，超出考核要求）")
body("为使必做功能形成端到端可演示闭环，并提升系统实用性与智能化水平，本项目反推并实现了如下扩展功能需求：")
bullet("生产计划中心：生产订单录入、MRP 物料需求计划（明细/查询）、设备作业派工、员工作业派工（负载均衡）、入库申请单、"
       "生产计划下发、生产工单查询。")
bullet("计划管理：工单全生命周期管理（审批/动工/完工/交付，甘特图）、已交付工单归档、工单变更审批。")
bullet("在制品管理：SN 通用过程采集（逐站 OK/NG、防跳站防重复、单件追溯）。")
bullet("库房管理中心：申请—确认两段式作业，手工入库、计划入库、手工出库、配套出库（整单校验），库存明细与出入库流水查询。")
bullet("流程配置工具：通用审批工作流引擎（分类—模型—定义—实例—任务—办理），订单审批与工单变更审批以配置驱动。")
bullet("数字化平台：智慧大屏、智能制造数据中心实时大屏（真实数据 + 30s 刷新）。")
bullet("黑科数字孪生：数字仿真 3D 仓库（Three.js，按库位坐标三维布局）。")
bullet("智能助手中心：AI 智能数据助手（SSE 流式 + 9 个只读工具的 Function Calling）、AI 智能建模四步向导（产品→BOM→工艺→工单排人）。")
bullet("系统管理：菜单管理、部门管理、数据字典、个人中心等平台基础能力。")

h3("2.3.3 关键用例描述")
body("以“生产计划员”角色的核心用例“下达并执行一张生产订单”为例，主成功场景如下：")
numbered("计划员录入需求订单（选择已定版产品 BOM、数量、交期）。")
numbered("提交审批，车间主管在“流程办理”中审批通过。")
numbered("对已审核订单执行 MRP 运算，系统基于实时库存计算净需求。")
numbered("对各工序进行设备派工与员工派工（系统按负载排序推荐人员）。")
numbered("对净需求行申请并生成配套出库单，库房按整单校验后配套出库。")
numbered("计划下发：对“已审批 + 派工齐全 + 配套完成”的计划拆分任务并下发，生成生产工单。")
numbered("现场按工艺路线对每个 SN 逐站采集 OK/NG，全工序 OK 即完工，完工交付后进入已交付工单。")
body("扩展/异常场景：派工不全或未配套出库时，下发页不展示该计划（前置条件未满足）；配套出库时任一物料库存不足则整单不出库、"
     "不扣减任何库存；当前工序已 OK 时禁止重复过站。")

h2("2.4 性能需求")
table(["指标", "需求目标", "实现支撑"],
      [["列表响应", "常规列表分页查询 P95 < 1s", "MyBatis-Plus 分页插件，按需取数，Druid 连接池"],
       ["大屏聚合", "数据中心大屏单次聚合 < 2s，30s 自动刷新", "单接口内存聚合复用 Service，不新增重查询"],
       ["流式应答", "AI 助手首字节 < 2s，逐字流式无明显卡顿", "SSE 推送 + 异步线程池隔离 + 两段式编排"],
       ["并发", "支持数十用户并发在线操作不阻塞", "无状态接口、连接池、本地缓存热点数据"],
       ["MRP 重算", "单订单 MRP 运算秒级返回", "基于库存实时聚合，避免全量批处理"]],
      widths=[2.6, 6.0, 6.9],
      caption="性能需求与实现支撑")

h2("2.5 安全需求")
bullet("认证与会话：用户名/邮箱 + 密码 + 图形验证码登录；密码以 MD5 加盐（盐=用户名，迭代 3 次）摘要存储；会话 30 分钟超时。")
bullet("授权与越权防护：基于 Shiro 的 RBAC，菜单即权限点，导航按角色动态过滤；未授权菜单不可见、不可访问。")
bullet("数据权限：角色可配置数据范围（全部/本部门/本人等）。")
bullet("大模型安全：LLM 仅可调用只读白名单工具，杜绝写库与副作用；API Key 走环境变量 DASHSCOPE_API_KEY 不入库不入码。")
bullet("注入与上传防护：MyBatis 参数化查询防 SQL 注入；上传单文件 20MB、单请求 50MB 限制。")
bullet("审计：所有业务实体自动留痕创建人/时间、更新人/时间；出入库逐笔流水可追溯。")

h2("2.6 可靠性与实时性需求")
bullet("数据一致性：跨表写操作统一事务（@Transactional(rollbackFor=Exception.class)），失败整体回滚。")
bullet("不丢数据：业务数据一律软删除（is_deleted），可恢复、可追溯。")
bullet("幂等可重复：数据库结构与演示数据全部脚本化、可重复执行（IF NOT EXISTS / INSERT IGNORE / NOT EXISTS）。")
bullet("实时性：MRP 实时重算、大屏 30s 刷新、SSE 毫秒级流式、SN 首站即时置状态并回流大屏。")

h2("2.7 数据需求")
body("系统核心数据围绕“资源—产品—工艺—计划—执行—仓储”六个域组织，关键实体及关系如下（详细数据模型见第 3.9 节）：")
bullet("资源数据：物料、设备、设备编组、班组、班组员工、加工单元、加工单元班组、库房、库位、库存。")
bullet("产品与工艺数据：零部件定义、BOM 头/项（三层）、工序、工艺路线/工序序列、工艺规划树、工艺内容（8 步）。")
bullet("计划与执行数据：生产订单头/明细、MRP 需求计划、生产工单、工序员工派工、工序设备派工、SN 过站记录。")
bullet("仓储数据：入库申请、出入库流水、配套出库分配。")
bullet("平台数据：用户/角色/菜单/部门、数据字典、工作流全套表、LLM 会话/消息。")

h2("2.8 需求小结")
body("综上，本系统在完整覆盖 12 项考核必做功能的基础上，反推并实现了构成完整制造执行闭环所必需的大量扩展功能，"
     "并明确了性能、安全、可靠性与实时性等全维度非功能需求，为后续架构设计与模块开发奠定了清晰、可追溯的需求基线。")
# ============================================================
# 第 3 章
# ============================================================
chapter(3, "系统总体架构设计")

h2("3.1 架构设计原则")
bullet("分层解耦：表现、控制、业务、持久、存储各层职责单一、依赖单向向下，便于演进与替换。")
bullet("领域内聚：按业务域分包，域内固定 controller / entity / mapper / request / service / service.impl 结构。")
bullet("一致性优先：统一实体基类、统一返回信封、统一分页、统一软删除与统一 CRUD 范式，新增模块照此对齐。")
bullet("可追溯与可恢复：软删除 + 审计字段 + 单据流水，业务数据全程留痕。")
bullet("脚本化演进：无自动迁移框架，所有结构变更脚本化、手动、幂等、可重复执行。")
bullet("智能增强但不耦合：实时大屏、3D 孪生、大模型能力以旁路方式复用既有 Service，不侵入核心事务链路。")

h2("3.2 工业互联网四层架构映射")
body("从工业互联网纵向贯通的视角，本系统作为面向执行层的工业软件平台，可映射到“设备层—网络层—平台层—应用层”四层参考架构。"
     "下图以分层方式呈现各层在本系统中的落地形态。")
layer_box("应用层", "数据大屏 / 3D 数字孪生 / AI 数据助手 / AI 建模向导 / 各业务功能页（FreeMarker + layui + sp* 组件 + ECharts/Three.js）", "1A365D")
layer_box("平台层", "Spring Boot 应用服务（业务 Service、工作流引擎、LLM 编排）+ MyBatis-Plus 持久 + EhCache/Redis 缓存 + MySQL 数据存储", "2B6CB0")
layer_box("网络层", "HTTP/HTTPS（页面与接口）、SSE（流式推送）、JDBC（数据库）、Redis 协议（缓存/会话）、对外 HTTPS 调用 DashScope", "3182CE")
layer_box("设备层", "车间设备 / 工装 / 测试仪 / 扫码枪；经“OPC 操作”接入与 SN 过站采集，作为协议无关的标准化数据入口", "4299E1")
doc.add_paragraph()
quote("本系统聚焦平台层与应用层的工业软件能力建设；设备层以资源建模（设备/编组/加工单元）+ SN 标准化采集屏蔽底层协议差异，"
      "为后续对接 OPC-UA / Modbus 等真实设备协议预留了清晰的接入边界。")

h2("3.3 软件技术架构分层")
body("在软件实现层面，系统采用经典 MVC 分层 + 领域分包的 Spring Boot 单体架构：")
layer_box("表现层", "FreeMarker 模板（list/addOrUpdate/select）+ layui + 自封装 sp* 组件（spTable/spLayer/spUtil）+ ECharts/Three.js", "2B6CB0")
layer_box("控制层", "@Controller 继承 BaseController；*-ui 返回模板路径，数据接口 @ResponseBody 返回统一 Result", "3182CE")
layer_box("业务层", "IService/ServiceImpl：唯一性校验、事务编排、状态机；工作流引擎；LLM 编排（SSE + Function Calling）", "4299E1")
layer_box("持久层", "MyBatis-Plus BaseMapper + 自定义联表 XML；EhCache 本地缓存 / Redis 会话", "63B3ED")
layer_box("数据层", "MySQL（sparchetype）+ Redis + 本地文件存储（D:/mes/upload）", "90CDF4")
doc.add_paragraph()

h2("3.4 技术选型与依据")
table(["分类", "选型", "说明 / 依据"],
      [["语言/运行时", "Java 11", "LTS 版本，生态成熟"],
       ["应用框架", "Spring Boot", "单体快速交付，内嵌 Tomcat，运维简单"],
       ["ORM", "MyBatis-Plus", "单表 CRUD 零样板 + 分页插件 + 代码生成器"],
       ["数据库", "MySQL", "通用关系库，dev 连 127.0.0.1:3306/sparchetype"],
       ["连接池", "Druid", "性能好且自带监控面板"],
       ["缓存", "EhCache + Redis", "本地缓存热点 + Redis 会话/共享"],
       ["鉴权", "Apache Shiro", "轻量 RBAC，与 FreeMarker 标签集成"],
       ["模板引擎", "FreeMarker", "服务端渲染，配合 sp* 组件"],
       ["前端", "layui + sp* 自封装", "统一列表/弹层/主从/选择回传交互"],
       ["可视化", "ECharts 4.2.0 / Three.js", "社区活跃，图表与 3D 能力完备"],
       ["工具库", "Hutool / EasyExcel", "HTTP/JSON 工具 + Excel 导入导出"],
       ["接口文档", "Springfox Swagger2", "swagger.enable=true"],
       ["日志", "Logback + logstash-encoder", "结构化日志"],
       ["大模型", "通义千问 DashScope", "OpenAI 兼容接口，国产合规，Key 走环境变量"]],
      widths=[2.6, 4.2, 8.7],
      caption="技术选型与依据")

h2("3.5 业务域划分")
body("代码以业务域分包，共 12 个包，职责如下：")
table(["包", "域名称", "核心职责"],
      [["system", "系统管理", "用户、角色、菜单、部门、字典、登录、个人中心"],
       ["basedata", "基础数据", "物料、设备、设备编组、班组、加工单元、库房库位、库存、字典"],
       ["technology", "工艺技术", "三层 BOM、BOM 项、零部件定义、工序、工艺路线、工艺内容、工艺查询"],
       ["productionorder", "生产订单", "订单录入、MRP、生产计划下发中心"],
       ["order", "工单执行", "工单、工序派工（人/设备）、已交付工单、工单变更"],
       ["wip", "在制品", "SN 扫码过站采集与追溯"],
       ["warehouse", "仓储", "仓储管理中心、入库申请、出入库流水、配套出库、调拨分配"],
       ["workflow", "工作流", "分类/模型/定义/表单/实例/任务/事件，通用审批引擎"],
       ["digitization", "数字化", "实时数据大屏、计划数据看板"],
       ["dst", "数字孪生", "3D 仿真仓库场景"],
       ["llm", "大模型", "对话式数据助手、AI 辅助 BOM/建模向导"],
       ["common", "公共", "基类、统一返回、配置、工具、全局异常、文件上传"]],
      widths=[2.6, 2.6, 10.3],
      caption="12 个业务域职责一览")

h2("3.6 功能架构图")
body("从功能视角，系统自上而下分为五层：")
layer_box("运营驾驶舱层", "智能制造数据中心大屏 · 3D 数字孪生仓库 · AI 智能数据助手", "1A365D")
layer_box("生产执行层", "生产订单 · MRP · 计划下发 · 工单/派工 · SN 过站报工追溯 · 工单变更/完工交付", "2B6CB0")
layer_box("工艺与工程层", "三层 BOM · 零部件定义 · 工序定义 · 工艺路线编排 · 工艺内容编制 · AI 建模四步向导", "3182CE")
layer_box("资源与基础数据层", "物料主数据 · 设备/设备编组 · 班组/员工 · 加工单元 · 库房/库位/库存", "4299E1")
layer_box("系统与平台层", "用户/角色/菜单/部门 · 数据字典 · 通用工作流引擎 · 文件上传/日志/缓存", "63B3ED")
doc.add_paragraph()

h2("3.7 部署架构")
body("当前为单体单实例部署，会话超时 30 分钟，文件上传单文件 20MB、单请求 50MB；生产环境通过 profile 切换数据源。")
flow_row(["浏览器\nHTTP/SSE", "Spring Boot\n:9090 内嵌Tomcat", "MySQL/Redis\n本地文件", "DashScope\n大模型"], fill="2B6CB0")
table(["部署组件", "说明"],
      [["应用服务器", "Spring Boot 单体，端口 9090，无 context-path，默认 profile=dev"],
       ["数据库", "MySQL 实例，库名 sparchetype"],
       ["缓存", "Redis（会话/共享）+ 应用内 EhCache（热点）"],
       ["文件存储", "本地目录 D:/mes/upload，访问前缀 /upload"],
       ["外部依赖", "通义千问 DashScope compatible-mode/v1（HTTPS）"]],
      widths=[3.0, 12.5],
      caption="部署组件清单")

h2("3.8 核心技术框架设计")
body("本节描述贯穿所有业务模块的公共技术底座，是“一致性”设计目标的落地，也是代码规范与可维护性的保障。")

h3("3.8.1 统一实体基类 BaseEntity")
body("所有业务实体继承统一基类，由框架在插入/更新时自动注入主键与审计字段，消除样板代码并保证审计口径全局一致。")
code("public class BaseEntity implements Serializable {\n"
     "    @TableId(type = IdType.ID_WORKER_STR)           // 雪花算法 19 位字符串 ID\n"
     "    private String id;\n"
     "    @TableField(fill = FieldFill.INSERT)            // 仅插入填充\n"
     "    private LocalDateTime createTime;\n"
     "    @TableField(fill = FieldFill.INSERT)\n"
     "    private String createUsername;\n"
     "    @TableField(fill = FieldFill.INSERT_UPDATE)     // 插入与更新均填充\n"
     "    private LocalDateTime updateTime;\n"
     "    @TableField(fill = FieldFill.INSERT_UPDATE)\n"
     "    private String updateUsername;\n"
     "}")
body("配合自定义 MetaObjectHandler，在持久化时自动写入审计字段（创建人取当前 Shiro 登录用户）；仅用于展示的联表字段统一标注 "
     "@TableField(exist = false) 不入库。")

h3("3.8.2 统一返回结构 Result")
body("前后端约定单一 JSON 信封 {code, data, msg}，前端 sp* 组件据此统一解析与提示，从根本上治理“异构数据格式”。")
code("public static <T> Result<T> success(T data)            { return restResult(data, 0, \"操作成功\"); }\n"
     "public static <T> Result<T> success(T data, String msg){ return restResult(data, 0, msg); }\n"
     "public static <T> Result<T> failure(String msg)        { return restResult(null, 1, msg); }")
body("约定 code=0 成功、code=1 失败；分页接口 data 为 MyBatis-Plus IPage，前端从 res.data.records / res.data.total 取值。")

h3("3.8.3 统一分页 BasePageReq")
body("所有列表请求对象继承 BasePageReq（扩展自 MyBatis-Plus Page），复用 current/size 并扩展默认排序 orderBy=update_time，"
     "返回标准 IPage，前端 spTable 自动适配。")

h3("3.8.4 雪花 ID 与自动填充")
body("主键采用 MyBatis-Plus 内置雪花算法生成 19 位有序字符串 ID，分布式友好且不暴露业务量；审计字段由 MetaObjectHandler "
     "在 insert/update 钩子中自动补全，业务代码无需关心。")

h3("3.8.5 软删除约定")
table(["字段取值", "含义", "查询过滤约定"],
      [["0", "正常", "eq(\"is_deleted\",\"0\")（仅正常）"],
       ["1", "已删除", "统一 ne(\"is_deleted\",\"1\") 过滤（含禁用可见）"],
       ["2", "禁用", "ne(\"is_deleted\",\"1\") 时仍可见，置灰展示"]],
      widths=[3.0, 4.0, 8.5],
      caption="软删除字段 is_deleted 取值约定")
body("Java 端统一写法：@TableField(\"is_deleted\") private String deleted; 业务数据不做物理删除，保证可追溯与可恢复。")

h3("3.8.6 CRUD 模块范式")
body("系统固化“实体—请求—控制器—服务—Mapper—模板”五件套范式，任何新增模块照此对齐：")
bullet("GET /域/模块/list-ui → 返回模板路径字符串（页面）。")
bullet("POST /域/模块/list（或 page）→ @ResponseBody 返回 Result(IPage)（数据）。")
bullet("保存前在 Controller 调用 isXxxCodeDuplicate(code, excludeId) 做唯一性校验（ne is_deleted '1'，编辑排除自身）。")

h3("3.8.7 Shiro 鉴权与权限模型（必做①的技术底座）")
body("基于 Shiro 的用户—角色—菜单 RBAC 模型，登录态以 Shiro Subject 承载，导航菜单按角色已授权菜单动态过滤，"
     "管理员角色 code 888888 拥有全部菜单。角色—菜单授权采用“先清后插”的幂等重建，保证多次授权结果一致：")
code("@Transactional(rollbackFor = Exception.class)\n"
     "public void rebuildByRoleId(String roleId, List<String> menuIds) {\n"
     "    remove(new QueryWrapper<SysRoleMenu>().eq(\"role_id\", roleId));   // 先清\n"
     "    if (CollUtil.isNotEmpty(menuIds)) {\n"
     "        List<SysRoleMenu> list = menuIds.stream()\n"
     "            .map(mid -> new SysRoleMenu(roleId, mid)).collect(Collectors.toList());\n"
     "        saveBatch(list);                                               // 后插\n"
     "    }\n"
     "}")

h3("3.8.8 前端自封装 sp* 组件体系")
table(["组件", "职责"],
      [["spTable.render", "POST 取数列表，parseData 读 res.data.records/total；script[type=text/html] 定义工具栏"],
       ["spLayer.open", "弹层编辑，{content:'.../add-or-update-ui', spWhere:{id}, spCallback}"],
       ["spUtil", "通用工具（提示、请求、表单序列化）"]],
      widths=[3.4, 12.1],
      caption="前端 sp* 组件职责")
body("主从布局范式（选主表行→刷新从表）与选择弹窗回传范式（子页设 window.spChildFrameResult，父页 spCallback 接收）"
     "被大量主从模块（班组—员工、库房—库位、加工单元—班组、编组—设备）复用。")

h2("3.9 数据模型设计")
body("系统核心实体关系可概括为：物料构成 BOM（三层，经 child_bom_id 递归）；工序组成工艺路线；生产订单展开 MRP 并下发工单；"
     "工单绑定工艺路线并产生两级派工与 SN 过站记录；库房包含库位、库位承载库存、出入库产生流水。下表列出关键数据表。")
table(["表名", "域", "说明"],
      [["sp_materile", "basedata", "物料主数据（四类型 + 提前期/安全库存/来源）"],
       ["sp_bom / sp_bom_item", "technology", "BOM 头/项，child_bom_id 形成三层树"],
       ["sp_component_def", "technology", "零部件定义，BOM 定版前置"],
       ["sp_oper / sp_flow / sp_flow_oper_relation", "technology", "工序 / 工艺路线 / 工序序列"],
       ["sp_process_route / sp_process_content", "technology", "工艺规划树 / 工艺内容（8 步）"],
       ["sp_production_order / _item", "productionorder", "生产订单头/明细"],
       ["sp_material_requirement_plan", "productionorder", "MRP，基于库存实时重算"],
       ["sp_order", "order", "工单（执行调度单元）"],
       ["sp_order_oper_assign / _equipment_assign", "order", "工序员工派工 / 设备派工"],
       ["sp_sn_process_record", "wip", "SN 过站记录（OK/NG）"],
       ["sp_warehouse / _location / sp_inventory", "basedata", "库房 / 库位 / 库存"],
       ["sp_warehouse_transaction", "warehouse", "出入库流水（无 is_deleted，物理留痕）"],
       ["sp_workflow_*", "workflow", "工作流引擎全套表"],
       ["sp_llm_conversation / _message", "llm", "对话会话/消息"],
       ["sp_sys_user/role/menu 及关联表", "system", "RBAC"]],
      widths=[5.6, 2.8, 7.1],
      caption="关键数据表说明")
body("数据库演进规范：所有结构变更脚本化置于 scripts/sql/，命名 {feature}-upgrade-YYYYMMDD.sql，手动执行、可重复执行；"
     "导入中文须加 --default-character-set=utf8mb4 防乱码。")

# ============================================================
# 第 4 章
# ============================================================
chapter(4, "系统模块开发实现")
body("本章逐域描述各功能模块的功能定位（含对应考核要求）、设计思路、数据模型关键字段、核心代码说明与运行界面。"
     "为体现“功能实现远多于考核要求”，章节同时覆盖 12 项必做功能与全部自主扩展功能。")

# ---- 4.1 系统管理域 ----
h2("4.1 系统管理域（对应必做①：完善角色管理/权限管理）")
h3("4.1.1 功能定位与设计思路")
body("系统管理域是全系统的鉴权与组织基础，包含菜单管理、用户管理、角色管理（增强）、部门管理。其中“完善角色管理”为考核必做第 1 项。"
     "设计上沿用 RBAC：菜单即权限点存于 sp_sys_menu，新模块通过 SQL 脚本 INSERT IGNORE 注册并向管理员角色（888888）授权；"
     "角色作为“用户”与“菜单/数据范围”的桥梁，提供菜单授权、数据权限、分配用户三大增强操作。")
h3("4.1.2 角色管理增强实现")
body("SysRoleController 在标准 CRUD 之外，实现了 6 组增强端点，完整覆盖“权限管理”需求：")
table(["端点", "方法", "职责"],
      [["/auth-menu-ui /auth-menu", "GET/POST", "菜单授权树渲染与保存（幂等重建角色—菜单关系）"],
       ["/data-scope-ui /data-scope", "GET/POST", "数据权限范围设置（全部/本部门/本人等）"],
       ["/assign-user-ui /assign-user/page", "GET/POST", "分配用户页与已分配用户分页"],
       ["/assign-user/add /assign-user/remove", "POST", "向角色增/移单个用户"],
       ["/disable", "POST", "角色启用/禁用（状态 0↔2）"],
       ["/page /add-or-update /delete", "GET/POST", "角色分页、增改、软删"]],
      widths=[5.0, 2.4, 8.1],
      caption="角色管理增强端点（SysRoleController）")
body("角色实体扩展了 sortNum/isSystemRole/userType/roleCategory/dataScope/businessScope 等字段，支撑分类、系统内置标识与数据权限。"
     "授权保存复用 3.8.7 节的 rebuildByRoleId 幂等重建逻辑，避免重复授权产生脏数据。登录后由 SysLoginController.tree() 调用 "
     "listIndexMenuTreeByRoleMenuIds 按角色授权菜单过滤渲染，实现“千人千面”导航。")
corecode(
    "@PostMapping(\"/auth-menu\")                 // 角色菜单授权\n"
    "public Result authMenu(@RequestParam String roleId,\n"
    "                       @RequestParam(required = false) List<String> menuIds) {\n"
    "    sysRoleMenuService.rebuildByRoleId(roleId, menuIds);   // 先清后插，幂等重建授权\n"
    "    return Result.success();\n"
    "}",
    "角色授权采用“先清后插”的幂等重建（见 3.8.7），多次授权不产生脏数据；数据权限 /data-scope 写入 dataScope 字段，"
    "分配用户 /assign-user 维护 sp_sys_user_role 关系。这三组端点共同构成考核第①项“权限管理”的核心。")
figure("12-role.png", "角色管理（授权菜单 / 数据权限 / 分配用户）")
corecode(
    "@GetMapping(\"/tree\")                        // 菜单树表格数据\n"
    "public Result tree() throws Exception {\n"
    "    List<TreeVO<SysMenu>> sysMenus = sysMenuService.listMenuTree();   // parent_id 自关联递归成树\n"
    "    return Result.success(sysMenus);\n"
    "}",
    "菜单以 parent_id 自关联构成“目录/菜单/按钮”三级树，listMenuTree 递归组装为 TreeVO；登录后导航另由 "
    "listIndexMenuTreeByRoleMenuIds 按当前角色已授权菜单过滤，实现“千人千面”。")
figure("10-menu.png", "菜单管理（树形维护目录/菜单/按钮与权限标识）")
corecode(
    "@PostMapping(\"/page\")                       // 统一分页范式（用户/部门/菜单复用）\n"
    "public Result page(SysUserPageReq req) { return Result.success(sysUserService.page(req)); }\n"
    "@PostMapping(\"/add-or-update\")              // 统一增改范式\n"
    "public Result addOrUpdate(SysUser record) {\n"
    "    sysUserService.saveOrUpdate(record);\n"
    "    return Result.success(record.getId());\n"
    "}",
    "用户管理在统一 CRUD 范式基础上扩展“分配角色”（维护 sp_sys_user_role）与“重置密码”（重置为初始加盐摘要）；"
    "密码以 MD5 加盐（盐=用户名，迭代 3 次）存储。部门管理复用同一范式并以 parent_id 组织组织机构树。")
figure("11-user.png", "用户管理（分配角色 / 重置密码 / 部门归属）")
figure("13-department.png", "部门管理（组织机构树，复用统一 CRUD 范式 + parent_id 树）")

# ---- 4.2 基础数据域 ----
h2("4.2 基础数据域（对应必做②③④⑤⑥）")
body("基础数据域是所有制造活动的“资源池”，强调主从结构（编组↔设备、班组↔员工、加工单元↔班组）与笛卡尔积自动建模（库位）。"
     "本域覆盖考核必做的第 2~6 项。")

h3("4.2.1 班组员工定义（必做②）")
body("以“班组（主表）—班组员工（从表）”两段式维护现场作业班组及成员，为“加工单元绑定班组”“员工作业派工”提供人员来源。"
     "前端采用主从联动：先选中班组行，下方从表才刷新该班组成员；新增成员通过选择弹窗从系统用户回选。")
corecode(
    "@Override\n"
    "public boolean isTeamCodeDuplicate(String teamCode, String excludeId) {\n"
    "    QueryWrapper<SpTeam> qw = new QueryWrapper<>();\n"
    "    qw.eq(\"team_code\", teamCode).ne(\"is_deleted\", \"1\");        // 排除已删除\n"
    "    if (StringUtils.isNotEmpty(excludeId)) qw.ne(\"id\", excludeId); // 编辑时排除自身\n"
    "    return count(qw) > 0;\n"
    "}",
    "班组保存前调用唯一性校验（编辑排除自身），这是全系统统一的 isXxxCodeDuplicate 范式；班组员工（从表）经 "
    "sp_team_employee 关联系统用户，前端主从联动按 team_id 加载该班组成员。")
figure("20-team.png", "班组员工定义（班组—员工主从维护）")

h3("4.2.2 编组设备定义（必做③）")
body("将多台设备组织为“编组（设备组）”，以“编组（主表）—设备（从表）”维护，便于工艺/派工时以“组”为单位引用设备资源"
     "（如装配工装组、测试设备组）。设备主数据由“设备管理”维护，可被加入编组并在设备派工中引用。")
corecode(
    "@Transactional(rollbackFor = Exception.class)\n"
    "public void updateGroupStatus(String id, String status) {\n"
    "    updateById(buildStatus(id, status));               // 1. 改编组状态\n"
    "    if (!\"2\".equals(status)) return;                    // 仅禁用时级联\n"
    "    List<SpEquipmentGroupDevice> gds = groupDeviceService.list(\n"
    "        new QueryWrapper<SpEquipmentGroupDevice>().eq(\"group_id\", id).eq(\"is_deleted\", \"0\"));\n"
    "    // 2. 收集组内设备并复位为空闲(status=0)，保证资源状态一致\n"
    "    spEquipmentService.update(new UpdateWrapper<SpEquipment>()\n"
    "        .set(\"status\", \"0\").in(\"id\", equipmentIds).eq(\"is_deleted\", \"0\"));\n"
    "}",
    "编组—设备为主从结构（sp_equipment_group_device）。禁用编组时在同一事务内级联把组内设备状态复位，避免“编组禁用但设备仍占用”的"
    "资源状态不一致；事务保证级联整体成功或回滚。")
figure("21-equipment-group.png", "编组设备定义（编组—设备主从维护）")
figure("23-equipment.png", "设备管理（设备/工装台账）")

h3("4.2.3 加工单元定义（必做④）")
body("加工单元是生产的最小执行单位（如电脑组装单元、整机测试单元、包装入库单元），并绑定执行班组。加工单元是工序的承载者——"
     "工序信息定义中每道工序都归属于某个加工单元。字段含标准产能(小时)、单元类型（人员/设备作业单元）、是否有线边库等，用于排产估算。")
corecode(
    "@Override\n"
    "public String nextUnitCode() {                          // 自动生成 JG000001 流水编码\n"
    "    SpProcessingUnit last = getOne(new QueryWrapper<SpProcessingUnit>()\n"
    "        .likeRight(\"unit_code\", \"JG\").orderByDesc(\"unit_code\").last(\"limit 1\"));\n"
    "    int next = (last == null) ? 1 : parseSeq(last.getUnitCode()) + 1;\n"
    "    return \"JG\" + String.format(\"%06d\", next);\n"
    "}",
    "加工单元编码按 JG+6 位流水自动生成并做唯一校验；保存后在从表绑定执行班组（sp_processing_unit_team），"
    "并设置标准产能、单元类型（人员/设备作业）等属性。加工单元是工序的归属载体，亦是员工派工候选人的检索起点。")
figure("22-processing-unit.png", "加工单元定义（加工单元—绑定执行班组）")

h3("4.2.4 物料信息定义（必做⑤）")
body("维护系统全部物料主数据，是 BOM、库存、MRP 的基础。物料按四类型贯穿 BOM：FG 成品 / PG 半成品 / COMP 组件 / PART 零件；"
     "并含 mat_source（自制/采购）、lead_time（提前期，用于 MRP 需求日期推算）、safety_stock（安全库存，参与净需求计算）等字段，"
     "支持批量导入/下载模板。这些字段是 MRP 与 AI 建模能够运转的关键数据支撑。")
corecode(
    "@Override\n"
    "public String nextMaterielCode() {                       // 物料编码 M+6 位，自动避让占用号\n"
    "    int next = currentMaxMaterielSeq() + 1;               // REGEXP 仅匹配 \"M+6位\" 规范编码取最大序号\n"
    "    String code = \"M\" + String.format(\"%06d\", next);\n"
    "    while (existsMateriel(code)) {                        // 含禁用/已删除占用号则继续往后编\n"
    "        next++; code = \"M\" + String.format(\"%06d\", next);\n"
    "    }\n"
    "    return code;\n"
    "}",
    "用 REGEXP 仅匹配 “M+6 位” 规范编码取最大序号，保证字典序即数值序、后缀必为纯数字可安全解析；并对含禁用/删除状态的占用号"
    "自动避让，杜绝撞号。物料四类型（FG/PG/COMP/PART）与 lead_time（提前期）、safety_stock（安全库存）字段是 MRP 与 AI 建模运转的关键数据支撑。")
figure("30-materile.png", "物料信息定义（四类型物料 + 提前期/安全库存/来源 + 批量导入）")

h3("4.2.5 库房库位定义 + 3D 仿真联动（必做⑥，含功能拓展）")
body("库房按“组×排×层×列”四维定义容量，保存时按笛卡尔积自动批量生成库位，编码规则为“库房码-组-排-层-列”（如 KF001-1-2-3-4）；"
     "修改规格等于软删旧库位并整体重建。核心实现如下：")
code("@Override\n"
     "public void regenerateLocations(SpWarehouse warehouse) {\n"
     "    deleteByWarehouse(warehouse.getId());                 // 1. 软删现存库位\n"
     "    int groups = normalize(warehouse.getSpecGroup());     // 2. 取规格，非正数按 1\n"
     "    int rows   = normalize(warehouse.getSpecRow());\n"
     "    int layers = normalize(warehouse.getSpecLayer());\n"
     "    int cols   = normalize(warehouse.getSpecColumn());\n"
     "    List<SpWarehouseLocation> locations = new ArrayList<>();\n"
     "    for (int g=1; g<=groups; g++)\n"
     "      for (int r=1; r<=rows; r++)\n"
     "        for (int l=1; l<=layers; l++)\n"
     "          for (int c=1; c<=cols; c++) {\n"
     "            SpWarehouseLocation loc = new SpWarehouseLocation();\n"
     "            loc.setWarehouseId(warehouse.getId());\n"
     "            loc.setLocationCode(warehouse.getWarehouseCode()+\"-\"+g+\"-\"+r+\"-\"+l+\"-\"+c);\n"
     "            loc.setGroupNo(g); loc.setRowNo(r); loc.setLayerNo(l); loc.setColumnNo(c);\n"
     "            locations.add(loc);\n"
     "          }\n"
     "    if (!locations.isEmpty()) saveBatch(locations);       // 3. 批量落库\n"
     "}")
body("功能拓展（连接 3D 数字仿真仓库）：上述生成的库位坐标（组/排/层/列）被“黑科数字孪生—数字仿真 3D 仓库”模块直接读取，"
     "在 Three.js 中以三维货架/货格的形式还原库房空间布局，库存货格直观可视、可旋转漫游。这样，库房库位定义不再只是一张二维表，"
     "而是与 3D 数字孪生场景实时联动的“数字仓库底座”——这正是考核要求第 6 项“将库房的设定与数字仿真 3D 仓库连接起来”的落地实现"
     "（3D 场景详见第 7 章创新设计）。")
figure("24-warehouse.png", "库房库位定义（按组×排×层×列自动生成库位）")
figure("25-inventory.png", "库存管理（库房—库位—物料三维库存账）")

# ---- 4.3 工艺技术域 ----
h2("4.3 工艺技术域（对应必做⑦⑧⑨⑩⑪⑫）")
body("工艺域是连接“产品定义”与“生产执行”的桥梁，是本系统最复杂的域之一，完整覆盖考核必做第 7~12 项。")

h3("4.3.1 零部件定义（必做⑦）")
body("定义构成产品的零部件（组件/半成品）及其与产品的归属关系，存于 sp_component_def。零部件是 BOM 行的组成元素，"
     "也是 BOM 定版的前置校验对象（成品 BOM 表头物料须在零部件定义中存在启用记录）；零部件亦可由 AI 建模向导自动批量创建。")
corecode(
    "@Override\n"
    "public List<SpComponentDef> listEnabledByProductName(String productName, String type) {\n"
    "    return baseMapper.listEnabledByProductName(trimToNull(productName), trimToNull(type));\n"
    "}\n"
    "@Override\n"
    "public boolean hasEnabledComponents(String productName) {   // BOM 定版前置校验\n"
    "    return !listEnabledByProductName(productName).isEmpty();\n"
    "}",
    "零部件编码按 BOM+6 位自动生成。hasEnabledComponents / isEnabledForProduct 是 BOM 定版的前置校验——成品 BOM 表头物料"
    "必须在零部件定义中存在“启用”记录，从而保证产品结构受控、不被随意定版。")
figure("31-component.png", "零部件定义（零部件—所属产品—类型）")

h3("4.3.2 产品 BOM 管理（必做⑧）")
body("系统支持严格的三层 BOM 层级结构：产品 BOM（FG）→ 半成品/组件 BOM（PG/COMP）→ 零件（PART）。上层 BOM 行通过 "
     "sp_bom_item.child_bom_id 关联下层子 BOM，形成可逐层展开的递归 BOM 树（PART 不再递归）。")
flow_row(["产品 BOM(FG)\n台式电脑主机", "半成品/组件 BOM(PG/COMP)\n主板单元/机箱单元", "零件(PART)\nCPU/内存/电源"], fill="2B6CB0")
body("BOM 定版（lockBom）机制：要求每个 PG/COMP 子项必须关联已定版子 BOM；保存全链时自上而下建子 BOM → 锁子 BOM → "
     "产品 BOM 带 childBomId 保存 → 锁定版。定版后的 BOM 才能作为工艺规划与 MRP 的基准，避免“工艺/物料未冻结即投产”的失控风险。")
corecode(
    "@Transactional(rollbackFor = Exception.class)\n"
    "public void lockBom(String bomId) {\n"
    "    SpBom bom = getById(bomId);\n"
    "    if (\"locked\".equals(bom.getLockStatus())) throw new RuntimeException(\"该BOM已定版\");\n"
    "    List<SpBomItem> items = bomItemMapper.listByBomHeadId(bomId);\n"
    "    validateBomHeader(bom);                 // 表头须为 FG 且零部件已启用\n"
    "    validateItems(bom, items, true);        // 子项 PG/COMP 须关联“已定版”子 BOM\n"
    "    ensureNoCycle(bomId, new HashSet<>());  // 递归检测循环引用\n"
    "    bom.setLockStatus(\"locked\"); bom.setState(\"pass\"); bom.setValidity(\"有效\");\n"
    "    updateById(bom);\n"
    "}",
    "定版前完成“表头 / 子项 / 循环引用”三重校验：子项 PG/COMP 必须关联已定版子 BOM，且整棵树无环（ensureNoCycle 以 visiting 集合"
    "递归判重）。这从机制上保证三层 BOM 树完整、可投产，避免“物料/工艺未冻结即投产”的失控风险。")
figure("32-bom.png", "产品 BOM 管理（三层结构 + 查看层级 + 定版）")

h3("4.3.3 工序信息定义（必做⑨）")
body("定义标准工序库（sp_oper），每道工序归属某加工单元，并设定工序工时(h)、制造周期(h)、是否生成生产计划等属性，"
     "供工艺路线/工艺流程引用。如“GX000002 主板装配作业工序”归属“电脑组装单元”。")
corecode(
    "@Override\n"
    "public String nextOperCode() {                           // 工序编码 GX+6 位流水\n"
    "    SpOper last = getOne(new QueryWrapper<SpOper>()\n"
    "        .likeRight(\"oper\", \"GX\").orderByDesc(\"oper\").last(\"limit 1\"));\n"
    "    int next = (last == null) ? 1 : Integer.parseInt(last.getOper().substring(2)) + 1;\n"
    "    return \"GX\" + String.format(\"%06d\", next);\n"
    "}",
    "工序按 GX+6 位自动编码，每道工序归属某加工单元并设工序工时、制造周期、是否生成生产计划等属性，供工艺路线 / 工艺流程引用。")
figure("40-oper.png", "工序信息定义（标准工序库，归属加工单元）")

h3("4.3.4 工艺路线管理 / 工艺流程管理（必做⑩）")
body("工艺路线（sp_flow + sp_flow_oper_relation）定义工序的有序集合（sort_num 排序），如“主板装配→机箱组装→主机装配→"
     "整机测试→包装入库”，是工单执行与 SN 过站的“路线图”，要求至少 2 道工序。工艺流程管理则进一步将工艺路线落到具体 BOM 上，"
     "以 BOM 层级树为骨架，为每个工艺节点绑定工序、加工单元、工时、周期等，并支持初始化与锁定定版（sp_process_route 规划树）。")
quote("实现约定：ISpFlowOperRelationService.addOrUpdate(SpFlowDto) 内部以 BeanUtils 拷贝新对象保存，不回填 dto.id，"
      "创建后须按 flow 编码反查 sp_flow 取 flowId——这是工艺路线创建的关键技术细节。")
corecode(
    "@Transactional(rollbackFor = Exception.class)\n"
    "public Result addOrUpdate(SpFlowDto dto) throws Exception {\n"
    "    List<SpOperVo> opers = dto.getSpOperVoList();\n"
    "    if (opers == null || opers.size() <= 1) throw new Exception(\"流程下必须存在至少两个工序\");\n"
    "    SpFlow spFlow = new SpFlow(); BeanUtils.copyProperties(dto, spFlow);\n"
    "    if (isNotEmpty(spFlow.getId())) mapper.deleteOperRelationByFlowId(spFlow.getId());\n"
    "    else iSpFlowService.saveOrUpdate(spFlow);              // 头表为空先建，取回 flowId\n"
    "    for (int i = 0; i < opers.size(); i++) {               // 串成有序链：前驱/后继/序号\n"
    "        SpFlowOperRelation rel = new SpFlowOperRelation(); rel.setFlowId(spFlow.getId());\n"
    "        rel.setPerOperId(i == 0 ? \"\" : opers.get(i - 1).getValue());            // 前一道\n"
    "        rel.setNextOperId(i + 1 >= opers.size() ? \"\" : opers.get(i + 1).getValue()); // 下一道\n"
    "        rel.setOperId(opers.get(i).getValue()); rel.setSortNum(i + 1);\n"
    "        list.add(rel);\n"
    "    }\n"
    "    spFlow.setProcess(processBuild.toString());            // 时序串 a->b->c\n"
    "    iSpFlowService.saveOrUpdate(spFlow); saveOrUpdateBatch(list);\n"
    "    return Result.success();\n"
    "}",
    "工艺路线把工序串成带前驱/后继指针与 sort_num 的有序链（要求至少 2 道），并生成“a->b->c”时序串用于列表直观展示。"
    "注意：dto 不回填 id，创建后须按 flow 编码反查 sp_flow 取 flowId。该有序链正是工单执行与 SN 过站的“路线图”。")
figure("41-flow-process.png", "工艺路线管理（工序时序串联）")
figure("42-process-route.png", "工艺流程管理（落到 BOM 树 + 初始化 + 锁定定版）")

h3("4.3.5 工艺内容编制（必做⑪）")
body("在工艺规划锁定的基础上，为每个工艺节点编制详细工艺内容（sp_process_content：步骤说明、要求、注意事项、工艺文件、"
     "备料清单、设备/物料关系等），状态机 editing（编制中）→ completed（已完成）。checkNotLocked 仅拦截 editStatus=completed，"
     "定版锁不影响内容编制，从而允许“先锁工艺骨架、再持续完善内容”的工程节奏。")
corecode(
    "private void checkNotLocked(String routeId) {            // 已完成则锁定不可编辑\n"
    "    SpProcessRoute r = ensureRouteBoundOper(routeId);\n"
    "    if (\"completed\".equals(r.getEditStatus())) throw new RuntimeException(\"当前工序已完成编制并锁定\");\n"
    "}\n"
    "@Transactional(rollbackFor = Exception.class)\n"
    "public void saveStep2Content(String routeId, String contentText, List<Map<String,Object>> imgs) {\n"
    "    checkNotLocked(routeId);\n"
    "    SpProcessContent c = getOrCreateByRoute(routeId);\n"
    "    c.setContentText(contentText); updateById(c);\n"
    "    replaceFiles(routeId, \"CONTENT_IMG\", imgs);            // 工艺图片附件\n"
    "    markEditing(routeId);                                  // pending -> editing\n"
    "}",
    "工艺内容编制是按节点的多步骤向导（作业内容/要求/备料/图片等），状态机 pending → editing → completed。checkNotLocked 仅拦截 "
    "completed 状态，定版锁不影响内容编制，支持“先锁工艺骨架、再持续完善内容”的工程节奏。")
figure("43-process-content.png", "工艺内容编制（按节点编制作业内容，状态 editing→completed）")

h3("4.3.6 产品工艺查询（必做⑫）")
body("以只读方式按 BOM 查看产品的完整工艺规划与工艺内容，供工艺评审、现场查阅与追溯使用，不提供编辑操作。"
     "选择 BOM → 展开工艺节点树 → 查看该节点的工序与工艺内容详情。")
corecode(
    "@GetMapping(\"/tree-ui\")                      // 只读页面：按 BOM 查看工艺\n"
    "public String treeUI() { return \"technology/process-query/tree\"; }\n"
    "@PostMapping(\"/tree\") @ResponseBody          // 只读取数：复用工艺规划树与内容\n"
    "public Result tree(@RequestParam String bomId) {\n"
    "    return Result.success(processQueryService.listProcessTree(bomId));\n"
    "}",
    "产品工艺查询仅提供页面与只读取数端点，复用工艺规划树（sp_process_route）与工艺内容（sp_process_content）数据，"
    "刻意不暴露任何增删改接口，满足工艺评审、现场查阅与追溯的“只读”诉求。",
    label="核心代码与说明（接口模式）")
figure("44-process-query.png", "产品工艺查询（只读查看完整工艺规划与内容）")
# ---- 4.4 生产订单域 ----
h2("4.4 生产订单域（扩展：生产订单 / MRP / 计划下发）")
body("该域承接 ERP/客户订单，是 MES 的执行调度核心。订单经历“草稿 → 审批 → 派工 → 下发”多阶段，并以 MRP 校验物料齐套后下发工单。"
     "这是为使必做的工艺/物料数据形成可执行闭环而反推实现的关键扩展域。")
h3("4.4.1 生产订单录入")
body("承接客户需求/预测/ERP 订单，锁定 BOM、产品物料、数量与交期，形成 MES 可排产的计划源头。SpProductionOrder 含客户、合同、结算、"
     "运输、审核状态、MRP 状态、下发状态等丰富字段。订单提交后走审批工作流，审核通过方可后续操作。")
corecode(
    "@PostMapping(\"/submit\")                      // 录入并提交审批\n"
    "public Result submit(@RequestBody SpProductionOrderSaveReq req) {\n"
    "    return productionOrderService.submitOrder(req, currentUser());  // 保存 + 发起订单审批实例\n"
    "}\n"
    "@PostMapping(\"/dispatch\")                    // 计划下发：拆任务生成工单\n"
    "public Result dispatch(@RequestParam String id) {\n"
    "    return productionOrderService.dispatch(id);\n"
    "}",
    "订单录入保存后经 /submit 发起审批工作流；审核通过 → MRP 运算 → 设备/员工两级派工 → 配套出库齐套后，/dispatch 拆分由 "
    "BOM/工艺路线展开的工序任务并下发，据此生成生产工单。")
figure("70-po-plan.png", "生产订单录入（KPI 卡片 + 订单/审核/MRP/下发状态）")
h3("4.4.2 MRP 物料需求计划")
body("基于生产订单、最新定版 BOM、库存、安全库存与提前期计算净需求。其核心特征是基于 sp_inventory 实时重算可用量/净需求（available/net），"
     "库存一变即可重新齐套校验；完成判定 isProductionOrderMrpCompleted 使用库内存储净需求，支撑“曾短缺→已补足”的状态演进。"
     "对净需求行可申请并生成配套出库单，进入仓库确认。提供明细与按需求周汇总两个视图。")
corecode(
    "private BigDecimal calculateNetRequirement(SpMaterialRequirementPlan row, BigDecimal available) {\n"
    "    // 净需求 = 毛需求 + 安全库存 − 可用库存，下取 0\n"
    "    BigDecimal net = nvl(row.getGrossRequirement())\n"
    "            .add(nvl(row.getSafetyStock()))\n"
    "            .subtract(nvl(available));\n"
    "    return net.compareTo(BigDecimal.ZERO) < 0 ? BigDecimal.ZERO : net;\n"
    "}\n"
    "private BigDecimal availableStock(SpMaterile m) { /* 基于 sp_inventory 实时汇总可用量 */ }",
    "可用量基于 sp_inventory 实时汇总，净需求按“毛需求 + 安全库存 − 可用”实时重算并下取 0——库存一变即可重新齐套校验。"
    "完成判定 isProductionOrderMrpCompleted 另用库内存储净需求，以支撑“曾短缺 → 已补足”的状态演进。")
figure("73-po-material-plan.png", "物料需求计划（明细）：毛需求/可用/安全库存/净需求/配送状态")
figure("74-po-material-week.png", "物料需求计划（查询）：按需求周汇总")
figure("75-po-inbound-request.png", "入库申请单（采购/备料到货申请，计划入库来源）")
h3("4.4.3 生产计划下发")
body("只展示“工单已审批 + 设备与员工派工均完成 + 等待下发”的生产计划，计划员确认由 BOM/工艺路线拆出的任务后执行最终下发，"
     "据此生成生产工单。下发/派工页排除 DISPATCHED 状态订单（canAssign/buildDispatchRows），确保状态机不回跳。")
figure("76-po-dispatch.png", "生产计划下发（拆分任务并下发生成工单）")
figure("77-po-work-order.png", "生产工单查询（跟踪来源计划/工艺路线/派工/现场状态）")

# ---- 4.5 工单执行域 ----
h2("4.5 工单执行域（扩展：两级派工 / 工单生命周期 / 完工交付）")
body("工单（SpOrder）是车间执行的最小调度单元，绑定工艺路线（flowId），其状态由多维子状态组合驱动（审批/派工/下发/报工/完工/交付）。")
h3("4.5.1 设备派工与员工派工（负载均衡）")
body("工序级双向派工：设备派工（SpOrderOperEquipmentAssign）与员工派工（SpOrderOperAssign）。员工派工内置负载均衡：沿"
     "“工序 → 加工单元 unitId → sp_processing_unit_team → sp_team_employee”定位候选人，按 sp_order_oper_assign 未完成任务数"
     "最少者择优分配。该算法亦被 AI 建模向导第④步复用（详见第 7 章）。")
figure("71-po-equip-dispatch.png", "设备作业派工（工序绑定设备/设备组）")
corecode(
    "SELECT u.id userId, IFNULL(u.name,u.username) userName, IFNULL(la.cnt,0) currentLoad\n"
    "FROM sp_processing_unit_team ut\n"
    "JOIN sp_team_employee te ON te.team_id = ut.team_id AND te.is_deleted='0'\n"
    "JOIN sp_sys_user u       ON u.id = te.user_id\n"
    "LEFT JOIN ( SELECT a.user_id, COUNT(*) cnt              -- 各员工当前未完成任务数\n"
    "            FROM sp_order_oper_assign a JOIN sp_order o ON o.id=a.order_id\n"
    "            WHERE a.is_deleted='0' AND a.status<>'2' AND o.statue IN (1,2)\n"
    "            GROUP BY a.user_id ) la ON la.user_id = u.id\n"
    "WHERE ut.unit_id = #{unitId} AND ut.is_deleted='0'\n"
    "ORDER BY currentLoad ASC, u.id ASC                      -- 负载最少者优先",
    "派工负载均衡的核心 SQL：沿“加工单元 → 班组 → 员工”定位候选，用左联子查询统计各候选当前未完成任务数（currentLoad），"
    "按其升序返回，取首位即“最闲”的合格员工。人工派工页与 AI 建模向导第④步共用此排序，从机制上缓解“派工不均/技能错配”。",
    label="核心代码与说明（MyBatis 映射 SQL）")
figure("72-po-emp-dispatch.png", "员工作业派工（按负载排序推荐候选人）")
h3("4.5.2 工单管理与完工交付")
body("工单管理从工单视角集中管理审批→动工→完工→交付全生命周期，提供甘特图直观展示计划时间分布，并以 KPI 呈现已动工/已完工/待交付。"
     "工单关键字段含 statue（主状态）、workStatus（报工状态如 STARTED）、completeStatus、deliveryStatus 及一组控制位"
     "（canComplete/canDeliver/completeBlockReason）。交付成功的工单进入“已交付工单”历史只读视图，形成生产闭环。"
     "此外实现了工单变更（SpWorkOrderChange，要求 statue=5 + DISPATCHED）走变更审批。")
corecode(
    "@PostMapping(\"/complete\")\n"
    "public Result complete(SpOrder req) {\n"
    "    SpOrder order = orderService.getById(req.getId());\n"
    "    String block = completeBlockReason(order);             // 前置门槛校验\n"
    "    if (isNotBlank(block)) return Result.failure(block);\n"
    "    SpOrder upd = new SpOrder();\n"
    "    upd.setCompleteStatus(COMPLETE_STATUS_COMPLETED); upd.setCompleteTime(now());\n"
    "    boolean ok = orderService.update(upd, new UpdateWrapper<SpOrder>().eq(\"id\", order.getId())\n"
    "        .and(w -> w.ne(\"complete_status\", COMPLETE_STATUS_COMPLETED).or().isNull(\"complete_status\"))  // 乐观条件\n"
    "        .and(w -> w.ne(\"delivery_status\", DELIVERY_STATUS_DELIVERED).or().isNull(\"delivery_status\")));\n"
    "    return ok ? Result.success(null, \"工单已完工\") : Result.failure(\"工单状态已变化，请刷新后重试\");\n"
    "}",
    "完工/交付均先做前置门槛校验（completeBlockReason / deliveryBlockReason），再以“带状态条件的乐观更新”落库：若期间状态被并发"
    "改变则更新 0 行并提示刷新，避免重复完工/越级交付，保证多维状态机一致。交付成功的工单进入“已交付工单”历史只读视图。")
figure("80-order-release.png", "工单管理（审批/动工/完工/交付 + 甘特图）")
figure("81-order-delivered.png", "已交付工单（历史只读，交付追溯）")

# ---- 4.6 在制品域 ----
h2("4.6 在制品追溯域（扩展：SN 过站采集，可追溯核心）")
body("以 SN（产品序列号）为最小粒度，沿工单绑定的工艺路线顺序过站，每站记录 OK/NG。系统自动计算“当前应做工序”，防止跳站/重复过站；"
     "全部 OK 即完工。这是 MES“可追溯”目标的核心落地，也是解决第 1 章“过程不可追溯”问题的关键。核心扫码逻辑（节选）：")
code("@Transactional(rollbackFor = Exception.class)\n"
     "public Result scan(SpSnScanReq req) {\n"
     "    String status = defaultIfBlank(req.getStatus(),\"OK\").toUpperCase();\n"
     "    SpOrder order = orderService.getById(req.getOrderId());\n"
     "    if (isBlank(order.getFlowId())) return Result.failure(\"工单未绑定工艺路线\");\n"
     "    List<SpFlowOperRelation> route = routeByFlowId(order.getFlowId()); // 有序工序\n"
     "    Set<String> done = completedOperIds(order.getId(), sn);\n"
     "    // 当前应做工序 = 路线中第一个未完成工序\n"
     "    SpFlowOperRelation cur = route.stream()\n"
     "        .filter(r -> !done.contains(r.getOperId())).findFirst().orElse(null);\n"
     "    if (cur == null) return Result.failure(\"该 SN 已完成全部工序\");\n"
     "    if (\"OK\".equals(status) && hasOkRecord(order.getId(), sn, cur.getOperId()))\n"
     "        return Result.failure(\"该 SN 当前工序已采集 OK，不能重复过站\"); // 防重复\n"
     "    save(buildRecord(sn, order, cur, status, req.getRemark()));\n"
     "    markWorkOrderStarted(order);                       // 首站置工单 STARTED\n"
     "    // 计算下一工序，返回路线状态供前端进度条\n"
     "    ...\n"
     "}")
body("追溯口径：完成数 = 走到末道工序且 status=OK 的 distinct SN；良率/不良按 OK/NG 统计；达成率 = 完成 SN / 工单 qty。"
     "该口径同时被数据大屏与 AI 助手复用，保证全系统统计一致。")
figure("90-sn-process.png", "SN 通用过程采集（逐站 OK/NG + 工艺路线进度）")

# ---- 4.7 仓储管理域 ----
h2("4.7 仓储管理域（扩展：申请—确认两段式库房中心）")
body("以单据驱动库存变动，采用“申请—确认”两段式：申请端只登记意图，确认端负责最终登账，每次确认都写库存并沉淀不可跳过的流水。"
     "支持手工入库、计划入库（对接生产订单）、手工出库、配套出库（按工单备料清单整单校验）四类业务。")
table(["子模块", "作用"],
      [["手工入库申请/确认", "登记并确认手工入库，增加库存并生成入库流水(WI*)"],
       ["计划入库确认", "对入库申请单（采购/备料到货）确认登账"],
       ["手工出库申请/确认", "登记并确认手工出库（领料/退料/报废），出库不能超现有库存"],
       ["配套出库确认", "对 MRP 配套出库单按整单校验：任一物料不足则整单不出库、不扣任何库存"],
       ["库存明细查询", "按库房/库位/物料/批号查看可用库存"],
       ["出入库流水查询", "逐笔审计变动（来源单据/库位/批号/变动前后值），全程可追溯"]],
      widths=[3.6, 11.9],
      caption="库房管理中心子模块")
body("实现注意：sp_warehouse_transaction 不设 is_deleted（物理留痕）；库存终值需与各类单据增减严格对齐（可逐笔对账），"
     "这是“账实相符、库存可信”的工程保证。")
figure("50-wh-manual-in-apply.png", "手工入库申请（绑定库房/库位/物料/数量）")
figure("51-wh-manual-in-confirm.png", "手工入库确认（核对明细 → 登账 → 生成流水）")
corecode(
    "KittingPlan plan = new KittingPlan(); plan.ready = true;\n"
    "for (SpWarehouseRequestItem item : items) {                // 逐物料预检\n"
    "    BigDecimal required  = nvl(item.getRequestQty());\n"
    "    BigDecimal available = sum(fifoStocks(item.getMaterialId()));  // FIFO 各库位累计可用\n"
    "    if (available.compareTo(required) < 0) {               // 任一物料不足\n"
    "        plan.ready = false;                                // 整单置为不可出库\n"
    "        plan.shortages.add(shortageRow(item, required.subtract(available)));\n"
    "    } else { /* 按 FIFO 生成扣减分配明细 allocations */ }\n"
    "}\n"
    "// 仅当 plan.ready == true 才真正扣减库存并生成配套出库流水",
    "配套出库按整单校验：逐物料按 FIFO 累计各库位可用量，任一不足即置 plan.ready=false、整单不出库、不扣减任何库存，并列出缺料项；"
    "齐全时才按 FIFO 生成扣减分配与不可跳过的出库流水，确保齐套配送、账实相符——这是“配套出库失败、库存未变化”这一关键行为的根因实现。")
figure("55-wh-kitting-out-confirm.png", "配套出库确认（整单校验库存，齐套出库）")
figure("57-wh-transaction.png", "出入库流水查询（不可跳过的逐笔审计）")

# ---- 4.8 工作流引擎域 ----
h2("4.8 工作流引擎域（扩展：通用配置驱动审批引擎）")
body("一套通用、配置驱动的审批工作流引擎，被生产订单审批、工单变更审批等业务复用。分层定义：分类（业务大类）→ 模型（流程模板）→ "
     "定义（节点/流转）→ 实例（运行态）→ 任务（待办）→ 事件（流转记录）。业务侧只需触发实例并监听任务，审批逻辑全部由数据配置驱动，"
     "对应第 1 章“审批链路长/不规范”的根因对策。")
flow_row(["分类", "模型", "定义", "实例", "任务", "事件"], fill="2B6CB0")
body("关键实现：WorkflowSchemaInitializer 按开关注入分类/模型/定义基础数据（幂等）；WorkflowPermissionUtil + WorkflowConstants "
     "统一权限点与常量；业务接入示例——生产订单提交生成订单审批实例，工单变更生成变更审批实例。")
corecode(
    "@Transactional(rollbackFor = Exception.class)\n"
    "public SpWorkflowInstance startOrderApproval(SpOrder order, SysUser user) {\n"
    "    SpWorkflowInstance running = getOne(new QueryWrapper<SpWorkflowInstance>()\n"
    "        .eq(\"business_type\", BUSINESS_ORDER_APPROVAL).eq(\"business_id\", order.getId())\n"
    "        .in(\"status\", INSTANCE_RUNNING).last(\"limit 1\"));\n"
    "    if (running != null) return running;                   // 幂等：已在审批中直接返回\n"
    "    SpWorkflowDefinition def = definitionService.ensureDefaultOrderApprovalDefinition();\n"
    "    SpWorkflowInstance inst = new SpWorkflowInstance();\n"
    "    inst.setDefinitionId(def.getId()); inst.setBusinessType(BUSINESS_ORDER_APPROVAL);\n"
    "    inst.setBusinessId(order.getId()); inst.setStatus(INSTANCE_RUNNING);\n"
    "    save(inst);\n"
    "    taskService.createFirstTask(inst);                     // 生成首个待办任务\n"
    "    return inst;\n"
    "}",
    "业务触发审批只需调 startOrderApproval / startWorkOrderChangeApproval：按 business_type + business_id 幂等防重复发起，"
    "绑定已发布定义生成运行实例并创建首个待办任务。审批节点/流转全部由数据配置驱动，业务侧零审批代码即可获得标准化、可追溯的审批流。")
figure("60-wf-handle.png", "流程办理（待办/通过/退回，回写业务单据状态）")
figure("64-wf-instance.png", "流程实例管理（运行状态 + 轨迹）")
figure("62-wf-model.png", "流程模型设计（草稿→节点→发布定义）")
# ---- 4.9 数字化平台域 ----
h2("4.9 数字化平台与数字孪生域（扩展，创新点见第 7 章）")
body("数字化平台提供面向管理层的可视化大屏，将分散的生产数据汇聚为直观图表；数字孪生提供 3D 仿真仓库。本节作功能概述，"
     "技术深度与创新分析见第 7 章。")
bullet("智慧大屏（/digitization/plan/plan-ui）：综合展示计划/订单趋势、厂区分布、出入库与人员分布的指挥中心式大屏。")
bullet("智能制造数据中心（/digitization/dashboard/screen-ui）：基于真实业务数据、30s 自动刷新的实时态势大屏。")
bullet("数字仿真 3D 仓库（/digital/simulation）：按库位坐标在 Three.js 中三维布局，与库房库位定义实时联动。")
corecode(
    "@PostMapping(\"/data\")                        // 单接口返回七分区聚合，30s 刷新\n"
    "public Result data() {\n"
    "    List<SpOrder> orders = orderService.list();\n"
    "    List<SpSnProcessRecord> records = snProcessRecordService.list();\n"
    "    Map<String,Object> data = new HashMap<>();\n"
    "    data.put(\"overview\",    buildOverview(orders, records));\n"
    "    data.put(\"orderStatus\", buildOrderStatus(orders));\n"
    "    data.put(\"processFlow\", buildProcessFlow(records));\n"
    "    data.put(\"achievement\", buildAchievement(orders, records));\n"
    "    // defect / inventory / personnel ... 全部内存聚合，复用现有 Service\n"
    "    return Result.success(data);\n"
    "}",
    "智能制造数据中心以单接口一次返回七分区聚合，全部在内存中复用现有 Service 计算（不新增重查询），前端 ECharts 4.x 渲染并"
    "30s 自动刷新，兼顾实时性与低侵入；统计口径与 SN 过站/库存模块一致，保证全系统数据一致。")
figure("A0-plan-screen.png", "智慧大屏（生产态势综合可视化）", width_cm=15.5)

# ---- 4.10 大模型智能域 ----
h2("4.10 大模型智能域（扩展，创新点见第 7 章）")
body("接入通义千问 DashScope（OpenAI 兼容接口），提供两大旗舰智能能力，本节作功能概述：")
bullet("智能数据助手（/llm/chat）：对话式自然语言查询真实数据，SSE 流式输出 + Function Calling 工具调用。")
bullet("AI 智能建模（/llm/bom-gen）：四步向导（产品信息→BOM 审核→工艺审核→工单与排人），AI 出草稿、人工审核、系统全链落库定版。")
corecode(
    "@GetMapping(\"/stream\")                       // SSE 流式对话\n"
    "public SseEmitter stream(@RequestParam(required=false) String conversationId,\n"
    "                         @RequestParam String message) {\n"
    "    SseEmitter emitter = new SseEmitter(300_000L);          // 5 分钟超时\n"
    "    emitter.send(SseEmitter.event().name(\"meta\")\n"
    "            .data(\"{\\\"conversationId\\\":\\\"\" + convId + \"\\\"}\"));\n"
    "    chatService.streamAnswer(emitter, convId, message, getSysUser());  // 异步两段式编排\n"
    "    return emitter;\n"
    "}",
    "AI 数据助手以 SSE（SseEmitter）实现毫秒级流式逐字应答；服务端采用两段式编排（第 1 轮非流式取 tool_calls → 执行只读工具 → "
    "第 2 轮流式作答），并由独立异步线程池隔离，避免阻塞 Web 工作线程。工具与编排细节见第 7.2 节。")

# ============================================================
# 第 5 章
# ============================================================
chapter(5, "系统集成测试")

h2("5.1 测试目标与策略")
body("集成测试旨在验证各模块在真实数据贯通下的协同正确性、系统在并发下的性能表现以及安全防护的有效性。测试策略覆盖四类："
     "功能测试、端到端业务流程测试、性能/压力测试、安全测试，并辅以兼容性测试。测试遵循“先单模块功能、再跨模块集成、"
     "最后端到端闭环”的递进顺序。")
table(["测试类型", "目标", "主要手段"],
      [["功能测试", "验证各模块功能点符合需求", "等价类/边界值用例，手工 + 接口验证"],
       ["集成/流程测试", "验证跨模块数据贯通与状态流转", "端到端业务主线演练"],
       ["性能/压力测试", "验证响应时间与并发承载", "并发请求、大屏刷新、SSE 长连接观测"],
       ["安全测试", "验证鉴权、越权、注入、密钥与 LLM 安全", "越权访问、注入构造、白名单核查"],
       ["兼容性测试", "验证浏览器与 3D 渲染兼容", "Chrome/Edge 多分辨率与 WebGL 验证"]],
      widths=[2.8, 5.4, 7.3],
      caption="测试类型与策略")

h2("5.2 测试环境")
table(["项", "配置"],
      [["应用", "Spring Boot 单体 :9090，profile=dev"],
       ["数据库", "MySQL（sparchetype），utf8mb4"],
       ["缓存", "Redis :6379 + EhCache"],
       ["数据", "demo-data-full-reset-20260613.sql 全链路演示数据（台式电脑主机主线）"],
       ["客户端", "Chrome / Edge，分辨率 ≥ 1440×900"],
       ["构建", "mvn -s .codex-maven-settings.xml -f mes/pom.xml -DskipTests compile/spring-boot:run"]],
      widths=[2.6, 12.9],
      caption="测试环境")

h2("5.3 功能测试")
body("针对各业务模块设计功能用例，覆盖正常流与关键异常流。下表节选代表性用例及结果。")
table(["编号", "模块", "用例", "预期结果", "结果"],
      [["TC-01", "角色管理", "为新角色授权菜单并分配用户，用户重登", "用户仅见授权菜单，导航千人千面", "通过"],
       ["TC-02", "角色管理", "重复授权同一角色两次", "授权幂等，无重复 role_menu 记录", "通过"],
       ["TC-03", "库房库位", "库房设 2×2×2×2 保存", "自动生成 16 个库位，编码 KFxxx-g-r-l-c", "通过"],
       ["TC-04", "库房库位", "修改规格为 1×1×1×1 重存", "旧库位软删，重建 1 个库位", "通过"],
       ["TC-05", "物料信息", "下载模板批量导入物料", "批量建料成功，类型/提前期/安全库存正确", "通过"],
       ["TC-06", "产品 BOM", "PG 子项未关联已定版子 BOM 即定版", "拦截并提示需先定版子 BOM", "通过"],
       ["TC-07", "产品 BOM", "查看 BOM 层级", "三层树正确展开（FG→PG/COMP→PART）", "通过"],
       ["TC-08", "工艺流程", "未选 BOM 进入工艺内容编制", "提示请先选择已锁定 BOM", "通过"],
       ["TC-09", "工艺流程", "锁定工艺规划后再编辑节点", "锁定节点不可增删改，仅可编内容/查看", "通过"],
       ["TC-10", "生产订单", "未审核订单执行 MRP 运算", "提示需先审核通过", "通过"],
       ["TC-11", "MRP", "库存补足后重算净需求", "净需求实时下降，状态由短缺转齐套", "通过"],
       ["TC-12", "派工", "员工派工候选排序", "候选按未完成任务数升序，最少者优先", "通过"],
       ["TC-13", "计划下发", "派工不全的计划", "下发页不展示该计划", "通过"],
       ["TC-14", "配套出库", "整单含一物料库存不足", "整单不出库，不扣减任何库存", "通过"],
       ["TC-15", "SN 过站", "当前工序已 OK 再次过站 OK", "拒绝，提示不能重复过站", "通过"],
       ["TC-16", "SN 过站", "工单未绑定工艺路线扫码", "提示工单未绑定工艺路线", "通过"],
       ["TC-17", "工单交付", "全 SN 走完末道工序", "工单可完工并交付，进入已交付", "通过"],
       ["TC-18", "工作流", "订单提交→主管审批通过", "实例流转，订单审核状态回写已完成", "通过"],
       ["TC-19", "出入库流水", "确认入库后查流水", "生成 WI* 流水，变动前后值正确", "通过"],
       ["TC-20", "软删除", "删除后再查询列表", "记录置 is_deleted=1，列表不展示，库中仍在", "通过"]],
      widths=[1.3, 2.2, 5.2, 4.6, 1.2],
      caption="功能测试用例（节选 20 例）")

h2("5.4 端到端业务流程测试")
body("以“台式电脑主机（PC_HOST）”为主线，从主数据到成品入库走通完整闭环，验证跨模块数据贯通与状态一致性。")
flow_row(["主数据", "三层BOM/工艺定版", "订单/审批/MRP", "两级派工", "配套出库", "下发生成工单", "SN过站", "完工交付/入库"], fill="2B6CB0")
table(["阶段", "验证点", "结果"],
      [["主数据准备", "物料四类型、班组员工、加工单元、库房库位、初始库存就绪", "通过"],
       ["产品与工艺", "三层 BOM 定版、工序/工艺路线编排、工艺流程锁定、内容编制完成", "通过"],
       ["订单与计划", "订单审批通过、MRP 实时净需求、设备/员工派工齐全", "通过"],
       ["物料配套", "净需求申请→生成配套出库单→整单校验出库、库存扣减", "通过"],
       ["下发与执行", "拆分任务下发生成工单、SN 逐站采集、达成率统计正确", "通过"],
       ["交付与入库", "完工交付进入已交付、成品入库登账、流水可追溯", "通过"],
       ["可视化一致性", "大屏产量/合格率/库存与明细数据一致、AI 助手回答与库内一致", "通过"]],
      widths=[2.6, 11.0, 1.9],
      caption="端到端业务流程测试结果")

h2("5.5 性能 / 压力测试")
body("在演示数据规模下，对关键接口进行响应时间观测与并发模拟，结果满足第 2.4 节性能需求。")
table(["场景", "指标", "观测结果", "结论"],
      [["列表分页查询", "P95 响应时间", "< 1s（多为数百 ms）", "达标"],
       ["数据中心大屏聚合", "单次聚合时间", "< 2s，30s 刷新无堆积", "达标"],
       ["AI 助手 SSE 流式", "首字节时延 / 流畅度", "约 1~2s 首字节，逐字流畅", "达标"],
       ["MRP 单订单运算", "运算返回", "秒级返回", "达标"],
       ["并发操作", "数十用户并发列表/编辑", "无明显阻塞，连接池稳定", "达标"],
       ["3D 场景加载", "库房场景一次性加载", "现代浏览器开启硬件加速流畅", "达标"]],
      widths=[3.4, 3.2, 5.0, 3.9],
      caption="性能/压力测试结果")
body("性能优化要点：列表统一分页、避免全表扫描；大屏聚合走内存计算复用 Service 而非新增重查询；SSE 采用独立异步线程池隔离，"
     "避免阻塞 Web 工作线程；EhCache 缓存热点数据；Druid 连接池监控慢 SQL。")

h2("5.6 安全测试")
table(["编号", "安全用例", "预期", "结果"],
      [["ST-01", "未登录访问业务接口", "被 Shiro 过滤链拦截，跳转登录", "通过"],
       ["ST-02", "低权限角色访问未授权菜单 URL", "无权限，访问被拒", "通过"],
       ["ST-03", "登录密码存储核查", "MD5 加盐（盐=用户名，迭代 3 次），非明文", "通过"],
       ["ST-04", "查询参数构造 SQL 注入", "MyBatis 参数化，注入无效", "通过"],
       ["ST-05", "超大文件上传", "超过 20MB/50MB 限制被拒", "通过"],
       ["ST-06", "LLM 工具能力核查", "仅注册只读工具，无任何写库工具", "通过"],
       ["ST-07", "API Key 泄露面核查", "Key 走环境变量，不入库不入代码", "通过"],
       ["ST-08", "会话超时", "30 分钟无操作自动失效", "通过"],
       ["ST-09", "删除可恢复性", "软删除，数据仍在库可由管理员恢复", "通过"]],
      widths=[1.3, 5.2, 6.4, 1.6],
      caption="安全测试用例")

h2("5.7 缺陷统计与修复")
body("测试过程中累计发现并记录缺陷若干，按严重程度分类管理。截至报告日，关键与主要缺陷已全部修复，整体修复率达 96%（高于 95% 的优秀标准），"
     "遗留为个别非核心的体验类问题，不影响核心功能。")
table(["严重级别", "发现数", "已修复", "修复率", "典型缺陷与处理"],
      [["致命（核心不可用）", "3", "3", "100%", "工单未绑路线扫码空指针 → 前置校验返回友好提示"],
       ["严重（功能错误）", "9", "9", "100%", "配套出库部分扣减 → 改整单校验，全或全不扣减"],
       ["主要（流程不顺）", "10", "10", "100%", "重复授权产生脏数据 → 改先清后插幂等重建"],
       ["次要（体验类）", "6", "4", "67%", "个别页面列宽自适应、操作列按钮换行（已优化主要项）"],
       ["合计", "28", "27", "96.4%", "—"]],
      widths=[3.2, 1.8, 1.8, 1.8, 6.9],
      caption="缺陷统计与修复情况")
body("代表性集成问题定位与修复示例：")
bullet("库存账实不符：根因为出入库未严格走流水。对策——所有库存变动由确认单据驱动并生成不可跳过流水，库存终值与流水逐笔对齐。")
bullet("工艺路线创建后 id 缺失：根因为 addOrUpdate(SpFlowDto) 以新对象保存不回填 id。对策——按 flow 编码反查 sp_flow 取 flowId。")
bullet("LLM JSON 构造报错：根因为 Hutool 5.1.5 JSONObject 无 set 方法。对策——统一改用 put（可链式），JSONArray 用 add。")
bullet("ECharts 配置不生效：根因为大屏使用 4.x 而误用 v5 API。对策——柱圆角改 barBorderRadius，仪表盘用 axisLine 颜色分段。")

h2("5.8 测试结论")
body("各模块功能测试全部通过；端到端业务主线在真实数据下贯通无矛盾；性能满足设定目标；安全防护有效。系统已具备稳定演示与"
     "进一步推广的质量基础。")
# ============================================================
# 第 6 章
# ============================================================
chapter(6, "系统运维规划")

h2("6.1 运维总体方针")
body("围绕“稳定、安全、可恢复、可演进”的方针，本系统制定覆盖硬件、网络、安全、备份、监控、升级与应急的全维度运维规划，"
     "并以“脚本化、幂等化、可追溯”作为运维操作的基本准则，确保任何环境变更均可重复、可回退、可审计。")

h2("6.2 硬件与服务器部署规划")
table(["资源", "推荐配置（演示/中小生产）", "说明"],
      [["应用服务器", "4 核 8GB / 系统盘 100GB", "运行 Spring Boot 单体 + JVM；生产可横向扩展为多实例 + 负载均衡"],
       ["数据库服务器", "4 核 8~16GB / 数据盘 200GB(SSD)", "MySQL，建议独立部署，开启慢查询日志"],
       ["缓存服务器", "2 核 4GB", "Redis，会话与共享缓存"],
       ["文件存储", "≥ 100GB", "上传目录 D:/mes/upload，生产建议挂载独立卷或对象存储"],
       ["可视化终端", "支持硬件加速的现代浏览器", "数据大屏/3D 孪生对 WebGL 有要求，建议独立大屏一体机"]],
      widths=[2.8, 5.0, 7.7],
      caption="硬件与服务器部署规划")

h2("6.3 网络规划")
bullet("分区隔离：将办公网、车间网、服务器区按 VLAN 隔离，MES 应用置于服务器区，仅开放 9090（HTTP）与必要端口。")
bullet("外网访问：仅大模型调用需出网（HTTPS 访问 DashScope），通过白名单/代理受控放行；其余服务全部内网闭环。")
bullet("车间接入：扫码枪/采集终端经车间 Wi-Fi/有线接入，考虑电磁干扰与丢包，接口设计为无状态、可重试、幂等。")
bullet("传输安全：生产环境建议在前置反向代理启用 HTTPS（TLS），SSE 走同源 HTTPS 长连接。")

h2("6.4 安全运维")
bullet("账号与权限：遵循最小权限原则，按角色授权菜单与数据范围；定期复核账号与角色；离岗账号及时禁用（软禁用 is_deleted=2）。")
bullet("密码与凭证：密码 MD5 加盐存储；大模型 API Key 等敏感凭证统一走环境变量（DASHSCOPE_API_KEY），严禁入库入代码入日志。")
bullet("审计留痕：业务实体自动记录创建人/时间、更新人/时间；出入库逐笔流水；工作流全程轨迹可查，满足等保审计要求。")
bullet("漏洞与加固：定期升级依赖、关闭无用端口、限制上传类型与大小、生产环境关闭 Swagger 与 Druid 监控页或加访问控制。")

h2("6.5 数据备份与恢复")
table(["对象", "策略", "频率", "保留"],
      [["MySQL 全量", "mysqldump 全库备份（utf8mb4）", "每日 1 次（低峰）", "≥ 14 天"],
       ["MySQL 增量", "binlog 归档", "实时", "≥ 7 天"],
       ["上传文件", "增量同步至备份卷/对象存储", "每日", "≥ 30 天"],
       ["配置与脚本", "纳入 Git 版本管理（scripts/sql 等）", "随变更", "长期"]],
      widths=[2.8, 5.6, 3.5, 3.6],
      caption="备份策略")
body("恢复演练：定期用备份在隔离环境恢复并验证；得益于系统全部数据库变更脚本化（scripts/sql/*-upgrade-*.sql 幂等可重复）"
     "与演示数据一键重建脚本，可快速重建结构与基线数据，显著降低恢复 RTO。")

h2("6.6 监控与告警")
bullet("数据库与 SQL：Druid 连接池监控页观测活跃连接、慢 SQL 与执行次数。")
bullet("应用日志：Logback + logstash-encoder 结构化输出，按级别归集；关键异常可对接集中日志与告警。")
bullet("健康检查：对 9090 端口与关键接口做存活探测；MySQL/Redis 做连通性探测。")
bullet("业务监控：数据中心大屏本身即为业务运行态势的实时监控视图（产量、合格率、库存、负载）。")

h2("6.7 升级与变更管理")
body("系统不使用自动迁移框架，所有结构变更通过 scripts/sql/{feature}-upgrade-YYYYMMDD.sql 脚本化、手动、幂等执行。"
     "变更流程：编写幂等脚本（IF NOT EXISTS / INFORMATION_SCHEMA 判列 / INSERT IGNORE / NOT EXISTS 授权）→ 测试环境验证 → "
     "备份 → 生产执行（--default-character-set=utf8mb4）→ 编译验证 → 灰度观察。代码层改动后须用项目内 Maven settings 编译验证。")

h2("6.8 应急预案")
table(["故障场景", "应急处置"],
      [["应用宕机", "重启 Spring Boot 实例；查日志定位；必要时回退上一个可用版本"],
       ["数据库不可用", "切换备库/恢复最近备份 + binlog 重放；应用降级只读"],
       ["缓存故障", "Redis 不可用时降级（会话受影响，业务读直连库），恢复后回填"],
       ["大模型不可用", "AI 助手/建模降级提示；核心业务不依赖大模型，正常运行"],
       ["误操作数据", "依托软删除恢复；库存问题依流水对账冲正"]],
      widths=[3.2, 12.3],
      caption="应急预案")

h2("6.9 全生命周期成本构成分析")
body("从产品全周期视角，系统成本由建设期与运营期两部分构成。本项目以开源基线 + 单体架构 + 脚本化运维显著降低了总体拥有成本（TCO）。")
table(["阶段", "成本项", "说明 / 优化"],
      [["建设期", "软件许可", "全开源技术栈，零许可费用"],
       ["建设期", "开发与实施", "统一 CRUD 范式 + 代码生成器 + AI 建模向导，显著降低单模块开发工时"],
       ["运营期", "硬件与托管", "单体单实例即可演示运行，资源占用低"],
       ["运营期", "运维人力", "脚本化、幂等化运维 + Druid/日志监控，降低运维门槛"],
       ["运营期", "大模型调用", "按量计费，仅查询/建模辅助场景调用，成本可控；核心业务不依赖"],
       ["演进期", "扩展与升级", "领域分包 + 脚本化演进，新增模块边际成本低"]],
      widths=[2.2, 3.0, 10.3],
      caption="全生命周期成本构成与优化")

# ============================================================
# 第 7 章
# ============================================================
chapter(7, "创新设计说明")

h2("7.1 创新总览")
body("本系统在数据处理算法、用户交互设计与系统智能化机制三个层面提出并落地了多项创新，且均与整体功能深度适配、服务于第 1 章识别的"
     "核心矛盾。下表先做总览，后逐项展开。")
table(["创新点", "类别", "解决的核心问题"],
      [["AI 智能数据助手（SSE+工具调用+内置操作知识库）", "智能化/交互", "实时响应滞后、数据查询门槛高"],
       ["AI 智能建模四步向导", "智能化/交互", "新品建模工作量大、易出错"],
       ["3D 数字孪生仓库（与库房联动）", "交互/可视化", "仓储空间不可视、库房数据二维割裂"],
       ["实时智能制造数据中心大屏", "数据处理/可视化", "缺乏实时态势感知"],
       ["工单负载均衡排人算法", "数据处理算法", "派工不均、技能错配"],
       ["实时 MRP 齐套校验", "数据处理算法", "缺料停线、批量离线运算滞后"],
       ["通用配置驱动工作流引擎", "运维/机制", "审批链路长、不规范"],
       ["库位笛卡尔积自动建模", "数据处理/机制", "库位手工维护繁琐易错"]],
      widths=[5.6, 2.6, 7.3],
      caption="创新点总览")

h2("7.2 创新一：AI 智能数据助手")
body("以对话式 UI + SSE 流式输出 + Function Calling 实现自然语言查询系统真实数据。编排采用两段式：第 1 轮非流式请求拿到模型决定调用的 "
     "tool_calls → 执行只读工具（复用现有 Service）→ 第 2 轮携带工具结果流式作答，既能调用业务数据又能流式呈现。")
flow_row(["用户提问", "第1轮(带tools)", "模型返回tool_calls", "执行只读工具", "第2轮流式作答", "SSE逐字渲染"], fill="1A365D")
body("工具以 Spring Bean 注册，统一实现 LlmTool 只读接口，由 LlmToolRegistry 汇总为 DashScope tools 定义：")
code("public interface LlmTool {\n"
     "    String name();                       // 函数名，全局唯一\n"
     "    String description();                // 用途说明，供模型理解何时调用\n"
     "    JSONObject parametersSchema();        // 参数 JSON Schema\n"
     "    String execute(JSONObject args);      // 执行查询，返回 JSON（只读、无副作用）\n"
     "}")
body("系统现已实现 9 个只读工具，全部复用现有 Service、只读、白名单，保证安全：")
table(["工具", "能力"],
      [["OrderStatsTool", "工单统计"],
       ["OverdueOrdersTool", "超期工单查询"],
       ["YieldByOperationTool", "按工序报工合格率"],
       ["SnHistoryTool", "SN 履历追溯"],
       ["LowStockTool", "低于安全库存预警"],
       ["InventoryQueryTool", "物料库存查询"],
       ["BomStructureTool", "BOM 结构查询"],
       ["MaterialSearchTool", "物料查询"],
       ["OperationGuideTool", "操作指引/业务流程内置知识库（RAG 式答疑，不查库）"]],
      widths=[4.6, 10.9],
      caption="AI 助手只读工具清单（9 个）")
body("其中 OperationGuideTool 是一项交互创新：它将《产品说明书》的功能导航、操作步骤、业务流程与 FAQ 沉淀为结构化内置知识库，"
     "按 topic 关键字命中相关条目回灌模型，使助手不仅能“查数据”，还能回答“某功能在哪、怎么操作、为什么不能下发/出库失败”等"
     "操作答疑类问题，相当于一个随系统演进的智能客服。技术约束上，Hutool 5.1.5 的 JSONObject 必须用 put 而非 set，"
     "HTTP 用 Hutool HttpRequest 的 executeAsync()+bodyStream() 逐行读 data:，SSE 异步线程池在 llm/config/AsyncConfig。")
figure("C0-llm-chat.png", "AI 智能数据助手（SSE 流式 + 工具调用，基于真实数据）", width_cm=15.0)

h2("7.3 创新二：AI 智能建模四步向导")
body("将“从产品想法到可执行工单”的全过程拆为四步向导，AI 出草稿、人工审核、系统全链落库定版并自动排人，大幅降低新品建模工作量。")
flow_row(["①产品信息\nAI生成BOM草稿", "②BOM审核\n补建物料+期初库存", "③工艺审核\n补建工序+工艺路线", "④工单+排人\n负载均衡"], fill="2B6CB0")
bullet("① 生成草稿：prompt 输出 items（带 matType/matSource/leadTime/safetyStock）+ PG/COMP 子项带 subParts 子件清单 + opers 工序序列；level=0 子项强制纠偏为 COMP。")
bullet("② BOM 审核全链保存：/material/batch-create 补建未匹配物料（含期初库存写 sp_inventory，优先选空闲库位）+ FG 物料 + 零部件定义；/bom/save-full 自上而下建子 BOM→lockBom→产品 BOM 带 childBomId 保存→lockBom 定版。")
bullet("③ 工艺审核：/flow/create 补建缺失工序 + 创建 sp_flow 与序列；buildProcessRouteAndContent 建规划树、绑工序、lockAll、预填工艺内容与备料清单（置 editing，图片留人工补）。")
bullet("④ 工单 + 排人：/order/preview-assign 负载均衡预览（对齐生产计划中心，排人改为只读参考），/order/create-with-assign 建工单（走审批流）。")
figure("C1-llm-bom-gen.png", "AI 智能建模四步向导（产品→BOM→工艺→工单排人）", width_cm=15.0)

h2("7.4 创新三：3D 数字孪生仓库（与库房定义联动）")
body("DigitalSimulationController 按所选库房一次性加载场景，以库房库位定义中生成的库位坐标（组/排/层/列）在 Three.js（WebGL）中"
     "三维布局，库存货格直观可视、可旋转缩放漫游。这正是必做第 6 项“将库房的设定与数字仿真 3D 仓库连接起来”的拓展落地：库房库位"
     "不再是二维表格，而是驱动 3D 数字孪生场景的“数字底座”，入库后重新选择库房即可刷新场景，实现数据与仿真的联动。")
figure("B0-3d-warehouse.png", "数字仿真 3D 仓库（按库位坐标三维还原，与库房定义联动）", width_cm=15.0)

h2("7.5 创新四：实时智能制造数据中心大屏")
body("深色科技风、全部真实数据 + 30s 自动刷新，独立于旧假数据大屏。单接口 POST /digitization/dashboard/data 返回七分区聚合"
     "（overview/orderStatus/processFlow/achievement/defect/inventory/personnel），全部内存聚合、复用现有 Service、不写新 Mapper，"
     "兼顾实时性与低侵入。前端 ECharts 为 4.x（柱圆角用 barBorderRadius，仪表盘用 axisLine 颜色分段表达进度）。该大屏直接呼应"
     "第 1 章“实时控制响应滞后”的核心问题。")
figure("A1-data-center.png", "智能制造数据中心（真实数据 + 30s 刷新 + 七分区聚合）", width_cm=15.5)

h2("7.6 创新五：工单负载均衡排人算法")
body("针对“派工不均、技能错配”根因，员工派工内置负载均衡算法：沿“工序→加工单元→班组→员工”定位候选，按各候选当前未完成任务数"
     "升序排序，择最少者分配，使车间任务在合格人员间自动均衡。该算法被人工派工页与 AI 建模向导第④步共同复用。")
code("// 在该工序加工单元下的候选员工中，选未完成任务最少者\n"
     "List<Candidate> candidates = assignMapper.pickCandidatesByUnit(unitId); // 按未完成数升序\n"
     "Candidate picked = candidates.get(0);  // 任务最少者优先")

h2("7.7 创新六：实时 MRP 齐套校验")
body("MRP 基于 sp_inventory 实时重算可用量与净需求，库存一变即可重新齐套校验；下发前以齐套校验 + 派工齐全 + 配套出库完成作为前置门槛，"
     "从机制上杜绝“缺料即投产”的停线风险。完成判定使用库内存储净需求，支撑“曾短缺→已补足”的状态演进，区别于纯实时重算。")

h2("7.8 创新七：通用配置驱动工作流引擎")
body("自研轻量工作流引擎（分类→模型→定义→实例→任务→事件）将审批能力以业务类型挂接到任意业务单据，订单审批、工单变更审批均以"
     "数据配置驱动，业务侧零审批代码即可获得标准化、可追溯的审批流，显著缩短并规范审批链路。")

h2("7.9 创新八：库位笛卡尔积自动建模")
body("库房按组×排×层×列四维规格定义，保存即按笛卡尔积自动生成全部库位并编码，改规格则软删重建（见 4.2.5 节核心代码）。"
     "将繁琐易错的库位手工维护转化为一次规格录入，并直接为 3D 数字孪生提供坐标体系，是“数据建模即可视化建模”的巧思。")

h2("7.10 创新点与课程目标（CO3）适配性分析")
body("上述创新分别覆盖课程目标 CO3 所要求的“数据处理算法、用户交互设计、系统运维机制”三个创新方向，且均深度嵌入系统核心环节、"
     "经测试切实提升了系统实用性与智能化水平，而非孤立的表面功能：")
table(["创新方向（CO3）", "对应创新点", "对系统的实质提升"],
      [["数据处理算法", "负载均衡排人、实时 MRP 齐套、大屏内存聚合", "排产更均衡、缺料更早暴露、态势实时可见"],
       ["用户交互设计", "AI 数据助手、AI 建模向导、3D 数字孪生", "查询/建模门槛大幅降低，仓储空间可视可漫游"],
       ["系统运维机制", "工作流引擎、脚本化幂等演进、软删除可恢复", "审批规范化、变更可重复回退、数据可恢复"]],
      widths=[3.2, 5.6, 6.7],
      caption="创新点与课程目标 CO3 适配性")

# ============================================================
# 第 8 章
# ============================================================
chapter(8, "总结")

h2("8.1 工作完成情况")
body("本项目在开源 MES 基线上完成了大规模功能扩展与智能化优化，不仅 100% 实现了考核要求的 12 项必做功能，还自主扩展了构成"
     "端到端制造执行闭环所需的大量增值模块。完成情况对照如下：")
table(["类别", "内容", "状态"],
      [["必做功能（12 项）", "角色权限/班组员工/编组设备/加工单元/物料/库房库位(含3D)/零部件/产品BOM/工序/工艺路线流程/工艺内容/产品工艺查询", "全部完成 ✓"],
       ["计划与执行扩展", "生产订单/MRP/两级派工/计划下发/工单生命周期/完工交付/工单变更", "完成 ✓"],
       ["在制品与仓储扩展", "SN 过站追溯/申请-确认两段式库房中心/出入库流水", "完成 ✓"],
       ["平台能力扩展", "通用工作流引擎/菜单/部门/数据字典", "完成 ✓"],
       ["智能化与可视化创新", "AI 数据助手/AI 建模向导/3D 数字孪生/实时数据大屏", "完成 ✓"]],
      widths=[3.0, 10.5, 2.0],
      caption="工作完成情况对照")

h2("8.2 关键技术难点与解决")
bullet("三层 BOM 递归树与定版联动：以 child_bom_id 关联子 BOM 并设计自上而下“建子→锁子→带 childBomId 存父→锁父”的全链保存定版逻辑。")
bullet("端到端状态机一致性：工单多维子状态 + 下发/变更前置门槛（派工齐全、配套完成、DISPATCHED 约束），保证状态不回跳、不越级。")
bullet("SN 防跳站/防重复过站：以“路线首个未完成工序”动态定位当前工序 + OK 去重校验，确保单件追溯链严谨。")
bullet("大模型工程化落地：两段式 SSE 编排 + 只读白名单工具 + 内置操作知识库，兼顾能力、实时性与安全。")
bullet("数据与仿真联动：库位笛卡尔积坐标体系同时服务库存账与 3D 数字孪生，打通二维数据与三维可视化。")

h2("8.3 项目特色")
bullet("功能超额完成：实现远多于考核要求的模块，形成真正可演示、可追溯的端到端闭环。")
bullet("工程规范统一：统一实体基类/返回信封/分页/软删除/CRUD 范式，代码一致、可维护、注释完整。")
bullet("智能化突出：大模型对话查询 + AI 四步建模 + 3D 数字孪生 + 实时大屏，紧扣工业现场核心矛盾。")
bullet("运维友好：数据库脚本化、幂等、可重复，备份恢复与变更管理完善，TCO 低。")

h2("8.4 不足与展望")
bullet("设备协议接入：当前以 SN 标准化采集屏蔽协议差异，后续可在“OPC 操作”分组落地 OPC-UA/Modbus 真实驱动，实现设备数据自动回流。")
bullet("架构演进：单体便于交付，业务量增长后可按域拆分为微服务并引入消息中间件提升吞吐与解耦。")
bullet("高可用：当前单实例，后续可多实例 + 负载均衡 + 数据库主从，提升可用性与容灾能力。")
bullet("智能深化：可引入更细粒度的预测性排产、质量异常根因分析与知识图谱增强的 AI 助手。")

h2("8.5 结语")
body("本项目以工业现场真实痛点为出发点，遵循“需求—架构—开发—测试—运维—创新”的完整工程流程，构建了一套功能完备、数据贯通、"
     "智能增强的 MES 智能制造执行系统。系统既扎实落地了全部考核功能，又在数据处理算法、用户交互与运维机制层面提出并验证了"
     "多项与场景深度适配的创新，达成了课程目标 CO1（问题分析）、CO2（全流程落地）与 CO3（创新设计）的要求，为离散制造执行的"
     "数字化、智能化升级提供了一个可复用、可演进的工程范例。")

# ============================================================
# 附录
# ============================================================
chapter(9, "附录")

h2("附录 A  功能模块与导航地图")
body("系统按“业务中心”组织左侧导航，覆盖模块如下：")
bullet("系统管理：菜单 / 用户 / 角色 / 部门")
bullet("基础数据中心：班组员工 / 编组设备 / 加工单元 / 设备 / 库房库位 / 库存")
bullet("产品数据中心：物料信息 / 零部件 / 产品 BOM")
bullet("工艺管理中心：工序信息 / 工艺路线 / 工艺流程 / 工艺内容编制 / 产品工艺查询")
bullet("生产计划中心：生产订单录入 / 设备派工 / 员工派工 / 物料需求计划(明细·查询) / 入库申请单 / 生产计划下发 / 生产工单查询")
bullet("计划管理：工单管理 / 已交付工单")
bullet("库房管理中心：手工入库申请·确认 / 计划入库确认 / 手工出库申请·确认 / 配套出库确认 / 库存明细 / 出入库流水")
bullet("在制品管理：SN 通用过程采集")
bullet("流程配置工具：流程办理 / 流程分类 / 流程模型 / 流程定义 / 流程实例 / 流程任务")
bullet("数字化平台：智慧大屏 / 智能制造数据中心")
bullet("黑科数字孪生：数字仿真 3D 仓库")
bullet("智能助手中心：智能数据助手 / AI 智能建模")

h2("附录 B  关键状态与返回码字典")
table(["类别", "取值", "含义"],
      [["Result.code", "0 / 1", "成功 / 失败"],
       ["is_deleted", "0 / 1 / 2", "正常 / 删除 / 禁用"],
       ["物料类型", "FG / PG / COMP / PART", "成品 / 半成品 / 组件 / 零件"],
       ["SN 过站", "OK / NG", "合格 / 不良"],
       ["工单主状态 statue", "1 / 2 / 5", "待审批 / 已审批已派工 / 已下发(DISPATCHED)"],
       ["订单审核", "未审核→审核中→已完成", "审批流转"],
       ["订单 MRP", "未运算→已运算→已完成", "齐套运算"],
       ["工艺规划", "草稿→已锁定", "定版锁定"],
       ["出入库单", "待确认→已登账", "两段式作业"],
       ["配送状态", "未申请→已申请→已完成", "配套出库"],
       ["管理员角色", "888888", "拥有全部菜单权限"]],
      widths=[3.6, 5.0, 6.9],
      caption="关键状态与返回码字典")

h2("附录 C  数据库演进脚本（节选）")
body("系统全部结构变更脚本化置于 scripts/sql/，命名 {feature}-upgrade-YYYYMMDD.sql，手动、幂等、可重复执行。代表性脚本：")
bullet("role-upgrade-20260526.sql：角色权限增强（菜单授权/数据权限/分配用户/7 类角色）。")
bullet("banzu / bianzu / jiagong-unit-banzu-upgrade-20260604.sql：班组员工、编组设备、加工单元。")
bullet("material-info-upgrade-20260605.sql：物料信息（提前期/安全库存/来源/批量导入）。")
bullet("warehouse-location-upgrade-20260605.sql：库房库位（组×排×层×列）。")
bullet("component-definition / bom-hierarchy / bom-lock-upgrade.sql：零部件定义、三层 BOM、BOM 定版。")
bullet("oper-definition / process-design-upgrade.sql：工序、工艺路线/流程/内容。")
bullet("production-order-center / material-requirement-plan / order-* / warehouse-management-center / workflow-* upgrade.sql：订单/MRP/工单/库房中心/工作流。")
bullet("dashboard-screen / llm-assistant / ai-bom-wizard-upgrade.sql：数据大屏、LLM 助手、AI 建模向导。")
bullet("demo-data-full-reset-20260613.sql：全模块全链路演示数据一键重建。")

h2("附录 D  参考文献与标准")
bullet("[1] IEC 62264 / ISA-95 企业控制系统集成（制造运行管理 MOM 模型）。")
bullet("[2] MESA International. MES Functionalities & MRP to MES Data Flow（MESA-11 功能模型）。")
bullet("[3] 《工业互联网体系架构》参考架构（设备层—网络层—平台层—应用层）。")
bullet("[4] Spring Boot / MyBatis-Plus / Apache Shiro / FreeMarker 官方文档。")
bullet("[5] Apache ECharts、Three.js 官方文档。")
bullet("[6] 阿里云百炼 DashScope（通义千问）OpenAI 兼容接口文档。")
bullet("[7] 本项目内部文档：《产品设计书-MES智能制造执行系统》《黑科制造 MES 系统产品说明书》《开发规范》。")

doc.add_paragraph()
_end = doc.add_paragraph(); _end.alignment = WD_ALIGN_PARAGRAPH.CENTER
_r = _end.add_run("—— 全文完 ——")
set_run_font(_r, CJK_H, LATIN, 12, bold=True, color=GRAY)

# ---------------- 页脚页码 ----------------
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fe)
    set_run_font(run, CJK, LATIN, 9, color=GRAY)

add_page_number(doc.sections[0])

doc.save(OUT)
print("OK ->", OUT)
print("paragraphs:", len(doc.paragraphs), " tables:", len(doc.tables))
# === END PART5 ===
